"""GUI smoke test — runs headless (offscreen).

    QT_QPA_PLATFORM=offscreen python tests/gui_smoke_test.py

Builds every page without opening a real window, runs a short measurement
session with the simulation instrument, produces a certificate, and exports
to Excel. Goal: catch runtime errors in the UI code (wrong Qt call, missing
signal, type error).
"""

import os
import sys
import tempfile

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

_tmp = tempfile.mkdtemp(prefix="callog-gui-")
from callog_common import db  # noqa: E402

db.DATA_DIR = _tmp
db.DB_PATH = os.path.join(_tmp, "test.db")

from callog_common import auth, certificate  # noqa: E402
from callog_common.qt import QT_BINDING, QtCore, QtGui, QtWidgets  # noqa: E402

PASSED, FAILED = [], []


def check(name, cond, detail=""):
    (PASSED if cond else FAILED).append(name)
    print("  [%s] %s  %s" % ("OK  " if cond else "HATA", name, detail if not cond else ""))


def pump(app, ms):
    """Spins the Qt event loop for ms milliseconds."""
    loop = QtCore.QEventLoop()
    QtCore.QTimer.singleShot(ms, loop.quit)
    loop.exec()


def _no_trend(dev, pump_fn, app):
    """Whether the trend sentence disappears from the note when the trend checkbox is unchecked."""
    dev.trend_chk.setChecked(False)
    pump_fn(app, 80)
    off = dev.drift_note.text()
    dev.trend_chk.setChecked(True)
    pump_fn(app, 80)
    return "/yıl" not in off and "r²" not in off


def main():
    print("\n=== CalLog arayuz duman testi ===")
    print("Qt baglamasi: %s\n" % QT_BINDING)

    db.connect()
    uid = auth.create_user("test", "Test Operator", "parola123", "admin")
    user = db.query_one("SELECT * FROM users WHERE id = ?", (uid,))

    app = QtWidgets.QApplication.instance() or QtWidgets.QApplication([])
    from callog_common import theme
    theme.apply(app, theme.LIGHT)

    # --- 1. Window and page setup --------------------------
    print("1. Arayuz kurulumu")
    from callog_common.ui.login import LoginDialog, NewUserDialog
    from callog_pulse_echo.ui.main_window import MainWindow

    win = MainWindow(user)
    check("ana pencere kuruldu", win is not None)
    check("sekiz sayfa yuklendi", win.tabs.count() == 8,
          "sayfa sayisi=%d" % win.tabs.count())
    check("olcum sekmesi baslangicta kapali",
          not win.tabs.isTabEnabled(3))

    win.home.refresh()
    check("ana ekran cihaz tablosu doldu", win.home.inst_table.rowCount() >= 2,
          "satir=%d" % win.home.inst_table.rowCount())

    dlg = LoginDialog()
    check("giris ekrani kuruldu", dlg is not None)
    nu = NewUserDialog(actor_id=uid)
    check("kullanici ekleme ekrani kuruldu", nu.role.count() == 3)

    win.history.reload()
    check("gecmis sayfasi yuklendi", win.history.table.rowCount() == 0)

    check("varsayilan tema beyaz", theme.current_mode() == theme.LIGHT)
    check("beyaz temada zemin acik",
          win.palette().color(QtGui.QPalette.Base).lightness() > 200)
    win._set_theme(theme.DARK)
    pump(app, 50)
    check("koyu temaya gecildi", theme.current_mode() == theme.DARK)
    check("koyu temada zemin koyu",
          win.palette().color(QtGui.QPalette.Base).lightness() < 128)
    win._set_theme(theme.LIGHT)
    pump(app, 50)
    check("beyaz temaya geri donuldu", theme.current_mode() == theme.LIGHT)

    # --- 2. Setup page --------------------------------------------
    print("\n2. Oturum hazirlik sayfasi")
    setup = win.setup
    check("cihaz listesi doldu", setup.instrument_combo.count() >= 2)

    sim_index = None
    for i in range(setup.instrument_combo.count()):
        if "SİMÜLASYON" in setup.instrument_combo.itemText(i):
            sim_index = i
            break
    check("simulasyon cihazi listede", sim_index is not None)
    setup.instrument_combo.setCurrentIndex(sim_index)
    pump(app, 50)
    check("fonksiyon listesi doldu (9 fonksiyon)",
          setup.function_combo.count() == 9,
          "adet=%d" % setup.function_combo.count())
    check("simulasyon uyarisi gosterildi",
          "filigranlı" in setup.status_label.text(), setup.status_label.text())

    setup.dut_company.setText("Ornek Devlet Hastanesi")
    setup.dut_manufacturer.setText("Fluke")
    setup.dut_model.setText("175")
    setup.dut_serial.setText("SN-GUI-001")
    setup.dut_type.setText("El tipi multimetre")
    setup.nominal_edit.setText("10.0")
    setup.tolerance_edit.setText("0.02")
    setup.env_temp.setText("23.2")
    setup.env_rh.setText("55.0")
    setup.env_pressure.setText("99.3")
    setup.interval_spin.setValue(0.1)
    setup.function_combo.setCurrentIndex(0)      # VDC

    started = {}
    setup.session_started.connect(
        lambda sid, drv: started.update({"sid": sid, "drv": drv}))
    setup._start()
    pump(app, 200)
    check("oturum baslatildi", "sid" in started, "sinyal gelmedi")
    # The signal connection from section 2 is still live, so started["sid"]
    # will be overwritten by later sessions -- keep the first session's id separately
    first_sid = started["sid"]

    # --- 3. Measurement page --------------------------------------------
    print("\n3. Canli olcum sayfasi")
    # begin() is not called: MainWindow is already connected to the
    # session_started signal, so this exercises the real flow.
    acq = win.acquire
    check("olcum sekmesine gecildi", win.tabs.currentWidget() is acq)
    check("olcum sekmesi etkinlesti", win.tabs.isTabEnabled(3))
    pump(app, 600)
    check("izleme modunda veri akiyor", len(acq._ys) > 3,
          "nokta=%d" % len(acq._ys))
    check("X ekseni saniye cinsinden artiyor",
          len(acq._xs) > 2 and acq._xs[-1] > acq._xs[0] and acq._xs[-1] < 60,
          "son=%.3f" % acq._xs[-1])
    check("izleme modunda kayit yok", acq.stats.n == 0)

    acq._start_recording()
    pump(app, 1500)
    check("kayit basladi ve okuma birikti", acq.stats.n > 10,
          "n=%d" % acq.stats.n)
    check("ortalama nominale yakin", abs(acq.stats.mean - 10.0) < 0.1,
          "mean=%.6f" % acq.stats.mean)
    check("istatistik paneli guncellendi",
          acq._stat_labels["n"].text() == str(acq.stats.n))
    check("min ve maks gosteriliyor",
          acq._stat_labels["min"].text() not in ("", "—")
          and acq._stat_labels["max"].text() not in ("", "—"),
          "min=%s maks=%s" % (acq._stat_labels["min"].text(),
                              acq._stat_labels["max"].text()))
    check("min <= maks", acq.stats.min <= acq.stats.max)

    acq.window_combo.setCurrentIndex(0)          # Last 10 sec
    acq.follow_chk.setChecked(True)
    pump(app, 300)
    (vx0, vx1), _ = acq.plot.getViewBox().viewRange()
    check("takip modu son veriyi izliyor", vx1 >= acq._last_t - 1.0,
          "gorunum=%.2f..%.2f son=%.2f" % (vx0, vx1, acq._last_t))
    acq.window_combo.setCurrentIndex(6)          # Keep zoom
    acq.plot.getViewBox().setXRange(acq._last_t - 3, acq._last_t, padding=0)
    pump(app, 400)
    (zx0, zx1), _ = acq.plot.getViewBox().viewRange()
    check("yakinlastirma sonrasi genislik korunuyor",
          abs((zx1 - zx0) - 3.0) < 1.5, "genislik=%.2f" % (zx1 - zx0))
    check("yakinlastirilmisken de kayiyor", zx1 >= acq._last_t - 1.0)
    acq._reset_view()
    pump(app, 100)
    check("uygunluk karari verildi",
          acq._stat_labels["verdict"].text() in ("UYGUN", "UYGUN DEĞİL"),
          acq._stat_labels["verdict"].text())

    # --- 3b. Stability indicator, outlier reading, tolerance warning ----
    print("\n3b. Kararlilik gostergesi ve uyarilar")
    from callog_common import stability

    check("kararlilik gostergesi dolu",
          acq.stability_label.text() not in ("", "—"),
          acq.stability_label.text())
    check("kararlilik durumu taninan bir deger",
          any(t in acq.stability_label.text()
              for t in stability.STATE_TR.values()),
          acq.stability_label.text())
    check("ilerleme cubugu okuma sayisini yaziyor",
          "okuma" in acq.progress.format(), acq.progress.format())
    # The simulation driver deliberately produces rare outliers, so by the
    # time execution reaches this line one may have already slipped through
    # and opened the alert bar (this was the race that made the test flaky).
    # Clear any accumulated alerts first so the rule is tested
    # deterministically: "no alert content means the bar is hidden".
    acq._clear_alert()
    check("uyari yokken serit gizli", acq.alert_label.isHidden())

    # A single reading far outside the band: both an outlier and out of
    # tolerance. Fed in manually since the simulation driver never produces
    # a value like this on its own.
    flagged_before = len(acq._flagged)
    band_before = acq._out_of_band
    acq._on_reading(99001, db.utc_now(), 25.0, "25.0", acq._last_t + 1.0)
    pump(app, 100)
    check("aykiri okuma isaretlendi", len(acq._flagged) > flagged_before,
          "%d -> %d" % (flagged_before, len(acq._flagged)))
    check("aykiri okuma tabloda AYKIRI yaziyor",
          any(acq.table.item(r, 4) and acq.table.item(r, 4).text() == "AYKIRI"
              for r in range(acq.table.rowCount())))
    check("tolerans disi okuma sayildi", acq._out_of_band > band_before)
    check("uyari seridi gorunur oldu",
          not acq.alert_label.isHidden() and acq.alert_label.text() != "")
    check("uyaride hem aykiri hem tolerans var",
          "aykırı" in acq.alert_label.text().lower()
          and "tolerans" in acq.alert_label.text().lower(),
          acq.alert_label.text())
    check("disla dugmesi acildi", not acq.flag_btn.isHidden())
    check("aykiri okuma yine de hesaba dahil", acq.stats.max >= 25.0,
          "maks=%s" % acq.stats.max)

    # Auto-stop: target reading count.
    # The target is captured first: the reading worker keeps running in the
    # background, so `stats.n` can grow between these two lines.
    target = acq.stats.n
    acq.target_spin.setValue(target)
    check("hedef verilince ilerleme cubugu olceklendi",
          acq.progress.maximum() == target,
          "maks=%d hedef=%d" % (acq.progress.maximum(), target))
    acq._maybe_autostop()
    check("hedefe ulasinca kayit durdu", not acq.recording)
    check("durunca buton yazisi degisti",
          acq.start_rec_btn.text() == "Kayıt durduruldu",
          acq.start_rec_btn.text())
    check("durdurma denetim kaydina gecti",
          db.query_one("SELECT COUNT(*) AS n FROM audit_log"
                       " WHERE action = 'session.recording_stop'")["n"] >= 1)
    check("oturum hala acik — otomatik durdurma oturumu kapatmiyor",
          db.query_one("SELECT status FROM sessions WHERE id = ?",
                       (started["sid"],))["status"] == "running")

    from callog_common.ui.acquire_page import export_plot
    png_path = os.path.join(_tmp, "olcum-grafik.png")
    export_plot(acq.plot, png_path)
    check("grafik PNG olarak yazildi",
          os.path.isfile(png_path) and os.path.getsize(png_path) > 1000,
          "%d bayt" % (os.path.getsize(png_path)
                       if os.path.isfile(png_path) else 0))

    # The new row must not become the widest thing on the page: the window
    # already barely fits on 1366px laptop screens, and every new check row
    # makes the problem worse.
    _cards = [c for c in acq.findChildren(QtWidgets.QFrame) if c.property("card")]
    _widths = [c.minimumSizeHint().width() for c in _cards]
    check("kararlilik seridi en genis satir degil",
          _widths[3] < max(_widths), str(_widths))
    check("olcum sayfasi asiri genislemedi",
          acq.minimumSizeHint().width() <= 1560,
          "min=%d" % acq.minimumSizeHint().width())

    acq.target_spin.setValue(0)
    acq.recording = True          # kalan testler kayit bekliyor
    acq.start_rec_btn.setText("Sertifikasyon sürüyor")

    acq.pause_btn.setChecked(True)
    pump(app, 200)
    n_paused = acq.stats.n
    pump(app, 400)
    check("duraklatinca okuma durdu", acq.stats.n == n_paused,
          "n %d -> %d" % (n_paused, acq.stats.n))
    acq.pause_btn.setChecked(False)
    pump(app, 300)

    acq.notes_edit.setPlainText("GUI duman testi oturumu.")
    acq.stop(status="completed")
    pump(app, 200)
    saved = db.query_one("SELECT COUNT(*) AS n FROM readings WHERE session_id = ?",
                         (started["sid"],))["n"]
    check("okumalar veritabanina yazildi (%d)" % saved, saved > 10)
    check("oturum tamamlandi olarak kapandi",
          db.query_one("SELECT status FROM sessions WHERE id = ?",
                       (started["sid"],))["status"] == "completed")

    # --- 4. History, certificate, Excel ---------------------------------
    print("\n4. Gecmis kayitlar ve ciktilar")
    win.history.reload()
    check("oturum gecmiste gorunuyor", win.history.table.rowCount() == 1)
    win.history.table.selectRow(0)
    pump(app, 100)
    check("detay paneli doldu", "Oturum #" in win.history.detail.toHtml())
    check("simulasyon oturumunda sertifika dugmesi acik",
          win.history.cert_btn.isEnabled())

    xlsx = os.path.join(_tmp, "cikti.xlsx")
    from callog_common.ui.history_page import _write_excel
    _write_excel(started["sid"], xlsx)
    check("Excel dosyasi uretildi", os.path.exists(xlsx)
          and os.path.getsize(xlsx) > 3000,
          "boyut=%d" % (os.path.getsize(xlsx) if os.path.exists(xlsx) else 0))

    # Watermarked certificate from a simulated session
    path, cert_no, result = certificate.build_pdf(started["sid"], uid)
    check("filigranli PDF uretildi (%s)" % cert_no,
          os.path.exists(path) and os.path.getsize(path) > 2000,
          "boyut=%d" % (os.path.getsize(path) if os.path.exists(path) else 0))
    check("SIM- serisinden numara aldi", cert_no.startswith("SIM-CAL-MED-"), cert_no)
    check("resmi numara serisi tuketilmedi",
          certificate.next_cert_no().endswith("0001"), certificate.next_cert_no())
    cert = db.query_one("SELECT * FROM certificates WHERE session_id = ?",
                        (started["sid"],))
    certificate.approve(cert["id"], uid)
    check("sertifika onaylandi",
          db.query_one("SELECT approved_at FROM certificates WHERE id = ?",
                       (cert["id"],))["approved_at"] is not None)

    # --- 5. Audit log --------------------------------------------
    print("\n5. Denetim kaydi")
    from callog_common import audit
    ok, bad, n = audit.verify_chain()
    check("hash zinciri saglam (%d kayit)" % n, ok, "bozuk satir=%s" % bad)
    actions = [r["action"] for r in db.query("SELECT action FROM audit_log ORDER BY id")]
    for expected in ("session.start", "session.recording_start", "session.completed",
                     "certificate.issue", "certificate.approve"):
        check("denetim kaydinda '%s' var" % expected, expected in actions)

    # --- 6. Calibrated device history -------------------------------
    print("\n6. Kalibre edilen cihaz gecmisi")
    setup._reload_dut_history()
    check("gecmis cihaz listesinde kayit var", setup.dut_table.rowCount() == 1,
          "satir=%d" % setup.dut_table.rowCount())

    setup.dut_table.selectRow(0)
    pump(app, 100)
    check("secilen cihazin son olcumu gosteriliyor",
          "Son ölçüm" in setup.dut_pick_label.text(), setup.dut_pick_label.text())

    for w in (setup.dut_company, setup.dut_manufacturer, setup.dut_model,
              setup.dut_serial, setup.nominal_edit):
        w.clear()
    setup._use_selected_dut()
    check("gecmis cihaz forma dolduruldu",
          setup.dut_serial.text() == "SN-GUI-001"
          and setup.dut_manufacturer.text() == "Fluke",
          "%s / %s" % (setup.dut_manufacturer.text(), setup.dut_serial.text()))
    check("onceki olcum ayarlari da getirildi",
          setup.nominal_edit.text() == "10", setup.nominal_edit.text())

    # The same serial number must not open a second DUT row
    before = db.query_one("SELECT COUNT(*) AS n FROM duts")["n"]
    setup._get_or_create_dut()
    after = db.query_one("SELECT COUNT(*) AS n FROM duts")["n"]
    check("ayni cihaz icin yeni kayit acilmiyor", before == after,
          "%d -> %d" % (before, after))

    # --- 7. Admin page --------------------------------------------
    print("\n7. Yonetim sayfasi")
    admin = win.admin
    admin.reload()
    check("yonetim sayfasi bes sekmeli", admin.tabs.count() == 5)
    check("kullanici tablosu doldu", admin.user_table.rowCount() >= 1)
    check("cihaz tablosu doldu", admin.inst_table.rowCount() >= 2)
    check("denetim kaydi tablosu doldu", admin.audit_table.rowCount() > 5)
    check("zincir dogrulama etiketi dolu",
          "sağlam" in admin.chain_label.text(), admin.chain_label.text())

    from callog_common import branding
    admin.org_name_edit.setText("Test Laboratuvarı")
    admin.department_edit.setText("Test Birimi")
    admin._save_branding()
    check("laboratuvar adi kaydedildi", branding.org_name() == "Test Laboratuvarı")
    check("birim adi kaydedildi", branding.department() == "Test Birimi")
    check("bas harfler hesaplaniyor", branding.initials() == "TL")

    op_id = auth.create_user("operator1", "Ikinci Operator", "parola123", "operator")
    auth.set_role(op_id, "approver", uid)
    check("rol degistirildi",
          db.query_one("SELECT role FROM users WHERE id = ?", (op_id,))["role"]
          == "approver")
    auth.set_active(op_id, False, uid)
    check("kullanici devre disi birakildi",
          db.query_one("SELECT is_active FROM users WHERE id = ?",
                       (op_id,))["is_active"] == 0)
    check("devre disi kullanici giris yapamiyor",
          auth.authenticate("operator1", "parola123") is None)
    try:
        auth.set_active(uid, False, uid)
        check("son yonetici korunuyor", False, "devre disi birakildi!")
    except ValueError as exc:
        check("son yonetici korunuyor", "tek yönetici" in str(exc), str(exc))

    # --- 8. Turkish character support ------------------------------------
    print("\n8. Turkce karakter destegi")
    from reportlab.pdfbase import pdfmetrics

    from callog_common import pdffont
    font_name, _bold, ascii_fallback = pdffont.register()
    check("PDF icin Unicode font bulundu (%s)" % font_name, not ascii_fallback,
          "Helvetica'ya dusuldu, Turkce karakterler bozulur")
    face = pdfmetrics.getFont(font_name).face
    missing = [ch for ch in pdffont.TR_CHARS if ord(ch) not in face.charToGlyph]
    check("Turkce karakterlerin hepsi fontta var", not missing,
          "eksik: %s" % "".join(missing))

    print("\n9. Hedef gosterimi, kriter, coklu dislama, filtreler, docx")

    # Auto-answer QMessageBox / QInputDialog prompts
    orig_question = QtWidgets.QMessageBox.question
    orig_information = QtWidgets.QMessageBox.information
    orig_get_text = QtWidgets.QInputDialog.getText
    QtWidgets.QMessageBox.question = staticmethod(
        lambda *a, **k: QtWidgets.QMessageBox.Yes)
    QtWidgets.QMessageBox.information = staticmethod(
        lambda *a, **k: QtWidgets.QMessageBox.Ok)
    QtWidgets.QInputDialog.getText = staticmethod(
        lambda *a, **k: ("Otomatik test gerekcesi", True))

    setup.dut_serial.setText("SN-GUI-002")
    setup.dut_model.setText("179")
    setup.nominal_edit.setText("10.0")
    setup.tolerance_edit.setText("-0.02")      # sign should be ignored
    idx = setup.criterion_combo.findData("minmax")
    setup.criterion_combo.setCurrentIndex(idx)
    setup.interval_spin.setValue(0.05)
    started2 = {}
    setup.session_started.connect(
        lambda sid, drv: started2.update({"sid": sid}))
    setup._start()
    pump(app, 300)
    sid2 = started2["sid"]
    row = db.query_one("SELECT tolerance, tolerance_mode FROM sessions WHERE id = ?",
                       (sid2,))
    check("tolerans mutlak deger olarak kaydedildi", row["tolerance"] == 0.02,
          str(row["tolerance"]))
    check("uygunluk kriteri kaydedildi", row["tolerance_mode"] == "minmax",
          row["tolerance_mode"])

    acq = win.acquire
    check("hedef satirinda nominal ve tolerans var",
          "10" in acq.target_label.text() and "0.02" in acq.target_label.text(),
          acq.target_label.text())
    check("hedef satirinda kriter yaziyor",
          "tüm okumalar" in acq.target_label.text().lower(), acq.target_label.text())

    check("baslamadan once buton 'Sertifikasyonu başlat'",
          acq.start_rec_btn.text() == "Sertifikasyonu başlat", acq.start_rec_btn.text())
    acq._start_recording()
    check("basildiktan sonra buton yazisi degisti",
          acq.start_rec_btn.text() == "Sertifikasyon sürüyor", acq.start_rec_btn.text())
    check("buton devre disi", not acq.start_rec_btn.isEnabled())
    pump(app, 1200)

    acq._flush()
    n_rows = acq.table.rowCount()
    check("tabloda birden fazla satir var", n_rows >= 4, "satir=%d" % n_rows)
    acq.table.clearSelection()
    for r in range(3):
        acq.table.selectionModel().select(
            acq.table.model().index(r, 0),
            QtCore.QItemSelectionModel.Select | QtCore.QItemSelectionModel.Rows)
    acq._exclude_selected()
    pump(app, 100)
    excluded = db.query_one(
        "SELECT COUNT(*) AS n FROM reading_exclusions e"
        " JOIN readings r ON r.id = e.reading_id WHERE r.session_id = ?",
        (sid2,))["n"]
    check("secili uc okuma birden dislandi", excluded == 3, "dislanan=%d" % excluded)

    # Right-click menu actions: undo exclusion, copy to clipboard
    acq._include_selected()
    pump(app, 100)
    left = db.query_one(
        "SELECT COUNT(*) AS n FROM reading_exclusions e"
        " JOIN readings r ON r.id = e.reading_id WHERE r.session_id = ?",
        (sid2,))["n"]
    check("dislama geri alindi", left == 0, "kalan=%d" % left)
    check("geri alma denetim kaydina gecti",
          db.query_one("SELECT COUNT(*) AS n FROM audit_log"
                       " WHERE action = 'reading.include'")["n"] == 1)
    check("geri alinan satir kayit durumuna dondu",
          acq.table.item(0, 4).text() == "kayıt", acq.table.item(0, 4).text())

    acq._copy_selected()
    clip = QtWidgets.QApplication.clipboard().text()
    check("okumalar panoya sekmeyle ayrilmis kopyalandi",
          clip.count("\n") == 3 and "\t" in clip and "Değer" in clip,
          repr(clip[:60]))

    acq.finish_btn.click()
    pump(app, 400)
    check("oturumu bitir dugmesi calisiyor",
          db.query_one("SELECT status FROM sessions WHERE id = ?",
                       (sid2,))["status"] == "completed")
    check("olcum sekmesi tekrar kapandi", not win.tabs.isTabEnabled(3))
    check("bitince buton yazisi geri dondu",
          win.acquire.start_rec_btn.text() == "Sertifikasyonu başlat")

    hist = win.history
    hist.reload()
    check("gecmiste iki oturum var", hist.table.rowCount() == 2,
          "satir=%d" % hist.table.rowCount())
    dut2 = db.query_one("SELECT dut_id FROM sessions WHERE id = ?", (sid2,))["dut_id"]
    i = hist.dut_filter.findData(dut2)
    check("cihaz filtresi acilirda listeleniyor", i >= 0)
    hist.dut_filter.setCurrentIndex(i)
    pump(app, 100)
    check("cihaz filtresi listeyi daraltti", hist.table.rowCount() == 1,
          "satir=%d" % hist.table.rowCount())
    hist.dut_filter.setCurrentIndex(0)
    pump(app, 100)

    docx_path = os.path.join(_tmp, "sertifika.docx")
    certificate.write_docx(started["sid"], docx_path)
    check("DOCX uretildi",
          os.path.exists(docx_path) and os.path.getsize(docx_path) > 5000,
          "boyut=%d" % (os.path.getsize(docx_path) if os.path.exists(docx_path) else 0))
    from docx import Document
    text = "\n".join(p.text for p in Document(docx_path).paragraphs)
    check("DOCX basligi Turkce", "KALİBRASYON SERTİFİKASI" in text)
    tables = Document(docx_path).tables
    cells = [c.text for t in tables for r in t.rows for c in r.cells]
    check("DOCX'te tolerans satiri var", any("Tolerans" in c for c in cells))
    check("DOCX'te uygunluk kriteri var", any("Uygunluk kriteri" in c for c in cells))

    # Certificate tab: filter and soft delete
    hist.tabs.setCurrentIndex(1)
    hist.reload_certificates()
    check("sertifika listesi doldu", hist.cert_table.rowCount() == 1,
          "satir=%d" % hist.cert_table.rowCount())
    hist.cert_state_filter.setCurrentIndex(2)          # Onaylanmış
    pump(app, 100)
    check("onaylanmis filtresi calisiyor", hist.cert_table.rowCount() == 1)
    hist.cert_state_filter.setCurrentIndex(1)          # Onay bekleyen
    pump(app, 100)
    check("onay bekleyen filtresi bosaltiyor", hist.cert_table.rowCount() == 0)
    hist.cert_state_filter.setCurrentIndex(0)
    pump(app, 100)

    hist.cert_table.selectRow(0)
    pump(app, 100)
    cert_row = hist._selected_certificate()
    check("silme dugmesi yonetici icin acik", hist.delete_btn.isEnabled())
    hist._delete_certificate()
    pump(app, 100)
    after = db.query_one("SELECT deleted_at, delete_reason FROM certificates"
                         " WHERE id = ?", (cert_row["id"],))
    check("sertifika silindi olarak isaretlendi", after["deleted_at"] is not None)
    check("silme gerekcesi kaydedildi", bool(after["delete_reason"]))
    check("kayit veritabaninda duruyor",
          db.query_one("SELECT COUNT(*) AS n FROM certificates")["n"] == 1)
    check("yonetici silinmis kaydi goruyor", hist.cert_table.rowCount() == 1,
          "satir=%d" % hist.cert_table.rowCount())

    # An operator must not see a deleted record
    hist.state.user = db.query_one(
        "SELECT * FROM users WHERE username = 'operator1'")
    hist.reload_certificates()
    check("operator silinmis kaydi gormuyor", hist.cert_table.rowCount() == 0,
          "satir=%d" % hist.cert_table.rowCount())
    hist.state.user = user
    hist.reload_certificates()
    hist.cert_table.selectRow(0)
    pump(app, 100)
    check("geri al dugmesi yoneticide acik", hist.restore_btn.isEnabled())
    hist._restore_certificate()
    pump(app, 100)
    check("sertifika geri alindi",
          db.query_one("SELECT deleted_at FROM certificates WHERE id = ?",
                       (cert_row["id"],))["deleted_at"] is None)

    # --- 10. Devices log ------------------------------------------
    print("\n10. Kalibre edilen cihazlar defteri")
    from callog_common import documents

    dev = win.devices
    dev.reload()
    check("cihaz defterinde iki cihaz var", dev.table.rowCount() == 2,
          "satir=%d" % dev.table.rowCount())

    dev.table.selectRow(0)
    pump(app, 200)
    dut_id = dev.selected_dut_id()
    check("cihaz secildi", dut_id is not None)
    check("ozet dolu", "Seri no" in dev.summary.toPlainText(),
          dev.summary.toPlainText()[:60])
    check("cihazin olcumleri listeleniyor", dev.session_table.rowCount() >= 1,
          "satir=%d" % dev.session_table.rowCount())

    dev.search.setText("SN-GUI-002")
    pump(app, 150)
    check("arama filtresi calisiyor", dev.table.rowCount() == 1,
          "satir=%d" % dev.table.rowCount())
    dev.search.clear()
    pump(app, 150)

    # Attach a legacy PDF report
    legacy = os.path.join(_tmp, "2023-eski-rapor.pdf")
    certificate.build_pdf.__module__       # (import check)
    with open(legacy, "wb") as fh:
        fh.write(b"%PDF-1.4\n% eski rapor testi\n")
    doc_id = documents.add(dut_id, legacy, "2023 kalibrasyon raporu",
                           "legacy_cert", uid, doc_date="2023-05-14",
                           notes="Uygulamadan onceki donem")
    check("belge eklendi", doc_id is not None)

    stored = documents.get(doc_id)
    check("dosya uygulama klasorune kopyalandi",
          os.path.exists(stored["file_path"])
          and os.path.abspath(stored["file_path"]) != os.path.abspath(legacy),
          stored["file_path"])
    check("sha256 kaydedildi", bool(stored["sha256"]))
    state, _msg = documents.verify(doc_id)
    check("dogrulama 'ok' donuyor", state == "ok", state)

    # The record survives even if the source file is deleted
    os.remove(legacy)
    state, _msg = documents.verify(doc_id)
    check("kaynak silinse de kopya duruyor", state == "ok", state)

    # A corrupted copy gets detected
    with open(stored["file_path"], "ab") as fh:
        fh.write(b"bozuldu")
    state, _msg = documents.verify(doc_id)
    check("dosya degisirse tespit ediliyor", state == "changed", state)

    dev._show_detail()
    check("belge listede gorunuyor", dev.doc_table.rowCount() == 1,
          "satir=%d" % dev.doc_table.rowCount())

    summary = documents.dut_summary(dut_id)
    check("ozet belge sayisini iceriyor", summary["counts"]["documents"] == 1)

    # Drift (trend) series
    series = documents.measurement_series(dut_id)
    check("olcum serisi uretildi", len(series) >= 1, "seri=%d" % len(series))
    dev._load_series(dut_id)
    check("seyir acilirinda nokta var", dev.series_combo.count() >= 1)
    check("seyir aciklamasi dolu", "Nominal" in dev.drift_note.text(),
          dev.drift_note.text()[:60])
    # Was the chart actually drawn: if left empty the view defaults to [0,1]
    (gx0, gx1), (gy0, gy1) = dev.drift_plot.getViewBox().viewRange()
    check("seyir grafigi cizildi (Y ekseni olcum degerine oturdu)",
          gy0 > 5 and gy1 < 15, "gorunum Y=%.3f..%.3f" % (gy0, gy1))
    check("olcum serisi tolerans da tasiyor",
          all(len(p) == 6 for pts in series.values() for p in pts))

    # The trend line can't be drawn without the same point spread across
    # multiple years. We fabricate three years of measurements and confirm
    # first the "at least 3 measurements" warning, then that a yearly rate
    # and limit-exceedance estimate get produced.
    check("uc olcumden az iken egilim uyarisi veriliyor",
          "en az" in dev.drift_note.text() or "/yıl" in dev.drift_note.text(),
          dev.drift_note.text()[-120:])

    import uuid as _uuid
    sim_inst = db.query_one(
        "SELECT id FROM instruments WHERE driver = 'simulated'")["id"]
    for year, value in ((2023, 10.00), (2024, 10.04), (2025, 10.08)):
        drift_sid = db.execute(
            "INSERT INTO sessions (uuid, name, operator_id, dut_id, instrument_id,"
            " function, unit, nominal, tolerance, tolerance_mode, started_at,"
            " ended_at, status, is_simulated)"
            " VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (str(_uuid.uuid4()), "Eğilim %d" % year, uid, dut_id, sim_inst,
             "VDC", "V", 10.0, 0.1, "mean", "%d-06-01T09:00:00+00:00" % year,
             "%d-06-01T09:10:00+00:00" % year, "completed", 1))
        for seq in range(5):
            db.execute(
                "INSERT INTO readings (session_id, seq, ts_utc, value, unit,"
                " elapsed_s) VALUES (?,?,?,?,?,?)",
                (drift_sid, seq, "%d-06-01T09:0%d:00+00:00" % (year, seq),
                 value + 0.0001 * seq, "V", float(seq)))

    dev._load_series(dut_id)
    notes = []
    for i in range(dev.series_combo.count()):
        dev.series_combo.setCurrentIndex(i)
        pump(app, 60)
        notes.append(dev.drift_note.text())
    trend_note = next((n for n in notes if "/yıl" in n), "")
    check("egilim yillik hiz olarak yaziliyor", bool(trend_note),
          " || ".join(n[-60:] for n in notes))
    check("egilim uyumu r2 ile veriliyor", "r²" in trend_note, trend_note[-160:])
    check("sinir asim tahmini uretildi",
          "sınırını" in trend_note or "on yıldan" in trend_note
          or "eğilim görünmüyor" in trend_note, trend_note[-200:])
    check("egilim kapatilabiliyor",
          dev.trend_chk.isChecked() and _no_trend(dev, pump, app))

    # Start a new measurement from the device page
    picked = {}
    dev.new_session_for.connect(lambda i: picked.update({"id": i}))
    dev._start_session()
    pump(app, 150)
    check("cihazdan yeni olcum sinyali gitti", picked.get("id") == dut_id)
    check("setup formu cihazla doldu",
          win.setup.dut_serial.text() == summary["dut"]["serial_no"],
          win.setup.dut_serial.text())
    check("yeni oturum sekmesine gecildi", win.tabs.currentIndex() == 2,
          "sekme=%d" % win.tabs.currentIndex())

    documents.remove(doc_id, uid, "Test temizligi")
    check("belge baglantisi kaldirildi", documents.get(doc_id) is None)
    check("dosya kopyasi diskte kaldi", os.path.exists(stored["file_path"]))

    # --- 10b. Add a new device (directly from the devices page) -----------
    print("\n10b. Yeni cihaz ekle")
    from callog_common.ui import devices_page

    # QDialog.exec() enters a real modal event loop, and nothing in
    # automation ever closes it, so it would wait forever. Same pattern used
    # elsewhere for QMessageBox.warning mocks: temporarily replace the
    # class's exec() with a function that fills in the form and returns.
    def _fill_and_accept(self):
        self._company.setText("Test Hastanesi")
        self._manufacturer.setText("Physio-Control")
        self._model.setText("LIFEPAK 15")
        self._serial.setText("SN-NEW-001")
        self._type.setText("Defibrilatör")
        return QtWidgets.QDialog.Accepted

    orig_new_dev_exec = devices_page._NewDeviceDialog.exec
    devices_page._NewDeviceDialog.exec = _fill_and_accept

    before_devices = dev.table.rowCount()
    dev._new_device()
    pump(app, 150)
    check("yeni cihaz listeye eklendi",
          dev.table.rowCount() == before_devices + 1,
          "satir=%d" % dev.table.rowCount())

    new_dut = db.query_one("SELECT * FROM duts WHERE serial_no = ?",
                           ("SN-NEW-001",))
    check("cihaz veritabanina yazildi", new_dut is not None)
    check("cihaz tipi kaydedildi",
          new_dut is not None and new_dut["device_type"] == "Defibrilatör")
    check("eklenen cihaz otomatik secildi",
          dev.selected_dut_id() == (new_dut["id"] if new_dut else None))
    check("cihaz eklemesi denetim kaydina gecti",
          db.query_one("SELECT COUNT(*) AS n FROM audit_log"
                       " WHERE action = 'dut.create' AND entity_id = ?",
                       (new_dut["id"],))["n"] == 1)

    # Resubmitting the same manufacturer + model + serial must not open a new row
    before_dupe = dev.table.rowCount()
    dev._new_device()
    pump(app, 150)
    check("ayni cihaz tekrar eklenmedi (mukerrer korumasi)",
          dev.table.rowCount() == before_dupe,
          "satir=%d" % dev.table.rowCount())

    # A record must not be created if a required field is left blank
    def _fill_blank_and_accept(self):
        return QtWidgets.QDialog.Accepted

    orig_warning2 = QtWidgets.QMessageBox.warning
    QtWidgets.QMessageBox.warning = staticmethod(
        lambda *a, **k: QtWidgets.QMessageBox.Ok)
    devices_page._NewDeviceDialog.exec = _fill_blank_and_accept
    before_blank = dev.table.rowCount()
    dev._new_device()
    pump(app, 100)
    check("bos zorunlu alanla cihaz eklenmedi",
          dev.table.rowCount() == before_blank,
          "satir=%d" % dev.table.rowCount())
    QtWidgets.QMessageBox.warning = orig_warning2

    devices_page._NewDeviceDialog.exec = orig_new_dev_exec

    # --- 11. Certificate chart ------------------------------------------
    print("\n11. Sertifikadaki olcum grafigi")
    from callog_common import chart, pdffont

    font_name, font_bold, _ = pdffont.register()
    inc, exc = chart.load_series(first_sid)
    check("grafik verisi okundu", len(inc) > 5, "dahil=%d dislanan=%d"
          % (len(inc), len(exc)))

    drawing = chart.session_drawing(first_sid, font=font_name,
                                    font_bold=font_bold)
    check("grafik cizimi uretildi", drawing is not None)
    check("cizimde eleman var", len(drawing.contents) > 20,
          "eleman=%d" % len(drawing.contents))

    chart_png = os.path.join(_tmp, "grafik.png")
    chart.png(first_sid, chart_png, scale=2.0, font=font_name,
              font_bold=font_bold)
    check("grafik PNG olarak yazildi",
          os.path.exists(chart_png) and os.path.getsize(chart_png) > 8000,
          "boyut=%d" % (os.path.getsize(chart_png)
                        if os.path.exists(chart_png) else 0))

    # The chart should be skipped for a session with no readings (the document should still be produced)
    empty_sid = db.execute(
        "INSERT INTO sessions (uuid, operator_id, dut_id, instrument_id, function,"
        " unit, started_at, status, is_simulated)"
        " VALUES (?,?,?,?,?,?,?,?,?)",
        ("bos-oturum", uid, dut2, db.query_one(
            "SELECT id FROM instruments LIMIT 1")["id"], "VDC", "V",
         db.utc_now(), "completed", 0))
    check("okumasiz oturumda grafik atlaniyor",
          chart.session_drawing(empty_sid, font=font_name) is None)

    # Did the PDF grow once it included the chart
    pdf_size = os.path.getsize(path)
    check("grafikli PDF makul boyutta (%d bayt)" % pdf_size,
          pdf_size > 15000, "grafik gomulmemis olabilir")

    # --- 12. Session naming and deletion ----------------------------------------
    print("\n12. Oturum adlandirma ve silme")
    from callog_common import sessions as session_svc

    row = db.query_one("SELECT * FROM sessions WHERE id = ?", (sid2,))
    check("oturuma varsayilan ad verildi", bool((row["name"] or "").strip()),
          repr(row["name"]))
    dut_row = db.query_one("SELECT company, serial_no FROM duts WHERE id = ?",
                           (row["dut_id"],))
    check("ad firma adini iceriyor", dut_row["company"] in row["name"], row["name"])
    check("ad seri numarasini iceriyor",
          dut_row["serial_no"] in row["name"], row["name"])
    check("ad tarih-saat iceriyor",
          row["started_at"][:10] in row["name"], row["name"])

    session_svc.rename(sid2, "Yillik kalibrasyon — 10 V noktasi", uid)
    check("yeniden adlandirma calisti",
          db.query_one("SELECT name FROM sessions WHERE id = ?",
                       (sid2,))["name"] == "Yillik kalibrasyon — 10 V noktasi")
    session_svc.rename(sid2, "", uid)
    check("bos ad varsayilana donuyor",
          dut_row["serial_no"] in db.query_one(
              "SELECT name FROM sessions WHERE id = ?", (sid2,))["name"])

    # A session that has a certificate cannot be deleted
    try:
        session_svc.soft_delete(first_sid, uid, "Deneme")
        check("sertifikali oturum korunuyor", False, "silindi!")
    except ValueError as exc:
        check("sertifikali oturum korunuyor", "sertifika" in str(exc).lower(),
              str(exc))

    # Deletion without a reason must be rejected
    try:
        session_svc.soft_delete(sid2, uid, "   ")
        check("gerekcesiz silme reddedildi", False, "silindi!")
    except ValueError as exc:
        check("gerekcesiz silme reddedildi", "gerekçe" in str(exc).lower(), str(exc))

    # Counts before deletion (section 11 also opened an extra empty session)
    hist.tabs.setCurrentIndex(0)
    hist.status_filter.setCurrentIndex(0)
    hist.reload_sessions()
    rows_before = hist.table.rowCount()
    sessions_before = documents.dut_summary(dut2)["counts"]["sessions"]

    session_svc.soft_delete(sid2, uid, "Yanlis nominal girilmis")
    deleted_row = db.query_one("SELECT * FROM sessions WHERE id = ?", (sid2,))
    check("oturum silindi olarak isaretlendi", deleted_row["deleted_at"] is not None)
    check("silme gerekcesi kaydedildi",
          deleted_row["delete_reason"] == "Yanlis nominal girilmis")
    check("okumalar veritabaninda duruyor",
          db.query_one("SELECT COUNT(*) AS n FROM readings WHERE session_id = ?",
                       (sid2,))["n"] > 0)

    hist.status_filter.setCurrentIndex(0)      # All statuses
    hist.reload_sessions()
    check("silinen oturum normal listede yok",
          hist.table.rowCount() == rows_before - 1,
          "%d -> %d" % (rows_before, hist.table.rowCount()))
    hist.status_filter.setCurrentIndex(4)      # Deleted
    hist.reload_sessions()
    check("yonetici silinmis oturumu goruyor", hist.table.rowCount() == 1,
          "satir=%d" % hist.table.rowCount())

    hist.state.user = db.query_one(
        "SELECT * FROM users WHERE username = 'operator1'")
    hist.reload_sessions()
    check("operator silinmis oturumu gormuyor", hist.table.rowCount() == 0,
          "satir=%d" % hist.table.rowCount())
    hist.state.user = user
    hist.reload_sessions()

    # The device log should not count a deleted session
    sessions_after = documents.dut_summary(dut2)["counts"]["sessions"]
    check("cihaz defteri silinen oturumu saymiyor",
          sessions_after == sessions_before - 1,
          "%d -> %d" % (sessions_before, sessions_after))

    session_svc.restore(sid2, uid)
    check("oturum geri alindi",
          db.query_one("SELECT deleted_at FROM sessions WHERE id = ?",
                       (sid2,))["deleted_at"] is None)

    QtWidgets.QMessageBox.question = orig_question
    QtWidgets.QMessageBox.information = orig_information
    QtWidgets.QInputDialog.getText = orig_get_text

    win.close()

    # --- 13. Role-based hiding -------------------------------------
    print("\n13. Rol tabanli gizleme")
    from callog_common import perms

    check("operator denetim kaydini goremiyor",
          not perms.can("operator", perms.VIEW_AUDIT))
    check("operator yonetim sayfasini goremiyor",
          not perms.can("operator", perms.VIEW_ADMIN))
    check("lab sorumlusu denetim kaydini goruyor",
          perms.can("approver", perms.VIEW_AUDIT))
    check("lab sorumlusu kullanici listesini goremiyor",
          not perms.can("approver", perms.VIEW_USERS))
    check("silinmis kayitlar yalnizca yoneticide",
          perms.can("admin", perms.SESSION_VIEW_DELETED)
          and not perms.can("approver", perms.SESSION_VIEW_DELETED))

    # A real window is built for each role and page/button visibility is
    # read back: the UI could still forget to consult the perms table even
    # if the table itself is correct.
    windows = {}
    for role in ("operator", "approver", "admin"):
        ruid = auth.create_user("rol_" + role, "Rol %s" % role, "parola123", role)
        rwin = MainWindow(db.query_one("SELECT * FROM users WHERE id = ?", (ruid,)))
        windows[role] = rwin

    op, ap, ad = windows["operator"], windows["approver"], windows["admin"]

    check("operatorde yonetim sayfasi hic yok",
          op.tabs.index_of("admin") < 0 and op.admin is None)
    check("operatorde alti sayfa var", op.tabs.count() == 6,
          "sayfa=%d" % op.tabs.count())
    check("lab sorumlusunda yonetim sayfasi var", ap.tabs.index_of("admin") >= 0)
    check("yoneticide sekiz sayfa var", ad.tabs.count() == 8,
          "sayfa=%d" % ad.tabs.count())
    check("onay kuyrugu yalnizca onaylayanlarda",
          op.tabs.index_of("approvals") < 0 and op.approvals is None
          and ap.tabs.index_of("approvals") >= 0
          and ad.tabs.index_of("approvals") >= 0)

    menus = lambda w: [m.title().replace("&", "") for m in w.menuBar().findChildren(
        QtWidgets.QMenu) if m.parent() is w.menuBar()]
    check("ses hizi sayfasi her rolde var",
          all(w.tabs.index_of("velocity") >= 0 for w in windows.values()))
    check("operator menusunde Yonetim yok", "Yönetim" not in menus(op),
          str(menus(op)))
    check("yonetici menusunde Yonetim var", "Yönetim" in menus(ad))

    check("lab sorumlusu kullanicilar bolumunu gormuyor",
          not ap.admin.has_section("users") and ap.admin.has_section("audit"))
    check("yonetici uc bolumu de goruyor",
          ad.admin.has_section("users") and ad.admin.has_section("audit")
          and ad.admin.has_section("instruments"))

    # isVisibleTo() can't be used here: the certificate buttons sit on the
    # inner tab's non-visible page, and since the ancestor is hidden it
    # always returns False. isHidden() reports the widget's own state directly.
    check("operatorde oturum silme dugmesi gizli",
          op.history.delete_session_btn.isHidden())
    check("lab sorumlusunda silme dugmesi acik",
          not ap.history.delete_session_btn.isHidden())
    check("geri al dugmesi yalnizca yoneticide",
          not ad.history.restore_session_btn.isHidden()
          and ap.history.restore_session_btn.isHidden())
    check("onay dugmesi operatorde gizli",
          op.history.approve_btn.isHidden()
          and not ap.history.approve_btn.isHidden())
    check("belge kaldirma dugmesi operatorde gizli",
          op.devices.remove_doc_btn.isHidden()
          and not ap.devices.remove_doc_btn.isHidden())
    check("silinmis durum suzgeci yalnizca yoneticide",
          ad.history.status_filter.count() == 5
          and op.history.status_filter.count() == 4,
          "yonetici=%d operator=%d" % (ad.history.status_filter.count(),
                                       op.history.status_filter.count()))

    # Service layer: must reject even if the UI is bypassed
    op_uid = db.query_one("SELECT id FROM users WHERE username = 'rol_operator'")["id"]
    try:
        session_svc.soft_delete(first_sid, op_uid, "yetkisiz deneme")
        denied = False
    except PermissionError:
        denied = True
    check("operator dogrudan cagriyla da silemiyor", denied)

    ap_uid = db.query_one("SELECT id FROM users WHERE username = 'rol_approver'")["id"]
    try:
        session_svc.restore(first_sid, ap_uid)
        restore_denied = False
    except (PermissionError, ValueError) as exc:
        restore_denied = isinstance(exc, PermissionError)
    check("lab sorumlusu geri alamiyor", restore_denied)

    for w in windows.values():
        w.close()

    # --- 13b. Approval queue, filters, backup, notifications -------------
    print("\n13b. Onay kuyrugu, tarih suzgeci, denetim aktarimi")
    from callog_common import backup, certificate as cert_svc, notifications, perms
    from callog_common.qt import Qt

    orig_question = QtWidgets.QMessageBox.question
    orig_information = QtWidgets.QMessageBox.information
    orig_get_text = QtWidgets.QInputDialog.getText
    orig_save = QtWidgets.QFileDialog.getSaveFileName
    QtWidgets.QMessageBox.question = staticmethod(
        lambda *a, **k: QtWidgets.QMessageBox.Yes)
    QtWidgets.QMessageBox.information = staticmethod(
        lambda *a, **k: QtWidgets.QMessageBox.Ok)
    QtWidgets.QInputDialog.getText = staticmethod(
        lambda *a, **k: ("Otomatik test gerekcesi", True))

    # Make sure the queue is non-empty: earlier sections may have approved
    # certificates or deleted their sessions. Approval and rejection are
    # tested separately, so two records are needed.
    while len(cert_svc.pending()) < 2:
        free = db.query_one(
            "SELECT s.id FROM sessions s"
            " LEFT JOIN certificates c ON c.session_id = s.id"
            " WHERE s.status = 'completed' AND s.deleted_at IS NULL"
            "   AND c.id IS NULL LIMIT 1")
        if free is not None:
            certificate.build_pdf(free["id"], uid)
            continue
        stale = db.query_one(
            "SELECT c.id FROM certificates c" + cert_svc.SOURCE_JOIN +
            " WHERE c.deleted_at IS NULL AND c.approved_at IS NOT NULL LIMIT 1")
        if stale is None:
            break
        db.execute("UPDATE certificates SET approved_at = NULL,"
                   " approved_by = NULL WHERE id = ?", (stale["id"],))

    queue = win.approvals
    check("onay kuyrugu sayfasi var", queue is not None)
    queue.reload()
    pending_rows = cert_svc.pending()
    check("kuyruk sorgusu ile tablo ortusuyor",
          queue.table.rowCount() == len(pending_rows),
          "tablo=%d sorgu=%d" % (queue.table.rowCount(), len(pending_rows)))
    check("kuyrukta onay bekleyen sertifika var", queue.table.rowCount() >= 1)
    check("kuyruk eskiden yeniye sirali",
          [r["issued_at"] for r in pending_rows]
          == sorted(r["issued_at"] for r in pending_rows))
    check("onaylanmis sertifika kuyrukta yok",
          all(r["approved_at"] is None for r in pending_rows))

    queue.table.selectRow(0)
    pump(app, 100)
    picked = queue._selected()
    check("kuyruktan sertifika secilebiliyor", picked is not None)
    check("secilen sertifikanin ozeti dolu",
          picked["cert_no"] in queue.summary.toPlainText(),
          queue.summary.toPlainText()[:80])
    check("okuma grafigi cizildi",
          "okuma" in queue.plot_note.text(), queue.plot_note.text())
    check("onayla dugmesi acildi", queue.approve_btn.isEnabled())

    queue._approve()
    pump(app, 100)
    check("kuyruktan onaylandi",
          db.query_one("SELECT approved_at FROM certificates WHERE id = ?",
                       (picked["id"],))["approved_at"] is not None)
    check("onaylanan kuyruktan dustu",
          all(r["id"] != picked["id"] for r in cert_svc.pending()))

    # Rejection: doesn't add a new state, just soft-deletes with a reason
    if queue.table.rowCount():
        queue.table.selectRow(0)
        rejected = queue._selected()
        queue._reject()
        pump(app, 100)
        row = db.query_one("SELECT deleted_at, delete_reason FROM certificates"
                           " WHERE id = ?", (rejected["id"],))
        check("geri cevrilen sertifika silinmis isaretlendi",
              row["deleted_at"] is not None)
        check("geri cevirme gerekcesi kaydedildi", bool(row["delete_reason"]))
        check("geri cevrilen kuyruktan dustu",
              all(r["id"] != rejected["id"] for r in cert_svc.pending()))

    # Date range filter — reset filters left over from earlier sections
    hist = win.history
    hist.search.setText("")
    hist.only_mine.setChecked(False)
    for combo in (hist.dut_filter, hist.instrument_filter, hist.status_filter,
                  hist.date_filter.combo):
        combo.setCurrentIndex(0)
    hist.reload()
    all_rows = hist.table.rowCount()
    # The count is variable since sessions were deleted in section 12; what
    # matters is that the filter can change this count.
    check("suzgecsiz oturum listesi dolu", all_rows >= 1,
          "satir=%d" % all_rows)
    idx = hist.date_filter.combo.findData(7)
    hist.date_filter.combo.setCurrentIndex(idx)
    pump(app, 100)
    # Section 10 generated sessions dated 2023-2025 for the trend test: the
    # last-7-days filter should exclude those but keep today's.
    recent_rows = hist.table.rowCount()
    check("son 7 gun suzgeci eski oturumlari ayikliyor",
          0 < recent_rows < all_rows, "%d / %d" % (recent_rows, all_rows))

    hist.date_filter.combo.setCurrentIndex(
        hist.date_filter.combo.findData("custom"))
    hist.date_filter.start.setDate(QtCore.QDate(2023, 1, 1))
    hist.date_filter.end.setDate(QtCore.QDate(2023, 12, 31))
    pump(app, 100)
    check("ozel aralik yalnizca o yilin oturumlarini getiriyor",
          hist.table.rowCount() == 1, "satir=%d" % hist.table.rowCount())

    hist.date_filter.combo.setCurrentIndex(
        hist.date_filter.combo.findData("custom"))
    old = QtCore.QDate(2000, 1, 1)
    hist.date_filter.start.setDate(old)
    hist.date_filter.end.setDate(QtCore.QDate(2000, 12, 31))
    pump(app, 100)
    check("gecmis bir araliga suzunce liste bosaliyor",
          hist.table.rowCount() == 0, "satir=%d" % hist.table.rowCount())
    check("ozel aralik alanlari gorunur oldu",
          not hist.date_filter.start.isHidden())
    check("aralik metni uretiliyor",
          hist.date_filter.describe() == "2000-01-01 – 2000-12-31",
          hist.date_filter.describe())
    hist.date_filter.start.setDate(QtCore.QDate(2000, 12, 31))
    hist.date_filter.end.setDate(old)
    check("ters girilen aralik duzeltiliyor",
          hist.date_filter.range() == ("2000-01-01", "2001-01-01"),
          str(hist.date_filter.range()))
    hist.date_filter.combo.setCurrentIndex(0)
    pump(app, 100)
    check("suzgec kaldirilinca liste geri geliyor",
          hist.table.rowCount() == all_rows)

    tab = hist.tabs.widget(0)
    before_w = tab.minimumSizeHint().width()
    hist.date_filter.setVisible(False)
    tab.layout().activate()
    pump(app, 50)
    check("tarih suzgeci sayfayi genisletmedi",
          tab.minimumSizeHint().width() == before_w,
          "%d -> %d" % (before_w, tab.minimumSizeHint().width()))
    hist.date_filter.setVisible(True)

    # The certificate tab has its own date filter
    hist.tabs.setCurrentIndex(1)
    hist.reload_certificates()
    cert_all = hist.cert_table.rowCount()
    hist.cert_date_filter.combo.setCurrentIndex(
        hist.cert_date_filter.combo.findData("custom"))
    hist.cert_date_filter.start.setDate(old)
    hist.cert_date_filter.end.setDate(QtCore.QDate(2000, 12, 31))
    pump(app, 100)
    check("sertifika listesi de tarihe gore suzuluyor",
          cert_all > 0 and hist.cert_table.rowCount() == 0,
          "%d -> %d" % (cert_all, hist.cert_table.rowCount()))
    hist.cert_date_filter.combo.setCurrentIndex(0)
    pump(app, 100)

    # Permission matrix
    admin_page = win.admin
    admin_page.reload()
    check("yetki matrisi sekmesi var", admin_page.has_section("perms"))
    check("yetki matrisi doldu",
          admin_page.perm_table.rowCount() == len(perms.matrix()),
          "satir=%d" % admin_page.perm_table.rowCount())
    check("matriste uc rol sutunu var",
          admin_page.perm_table.columnCount() == 2 + len(perms.ROLE_ORDER))
    marks = {admin_page.perm_table.item(r, c).text()
             for r in range(admin_page.perm_table.rowCount())
             for c in range(2, admin_page.perm_table.columnCount())}
    check("matris hucreleri yalnizca isaret tasiyor", marks <= {"✓", "—"},
          str(marks))

    # Audit log export
    csv_path = os.path.join(_tmp, "denetim.csv")
    QtWidgets.QFileDialog.getSaveFileName = staticmethod(
        lambda *a, **k: (csv_path, "CSV (*.csv)"))
    admin_page.audit_search.setText("")
    admin_page.audit_action.setCurrentIndex(0)
    admin_page.audit_dates.combo.setCurrentIndex(0)
    exported = len(admin_page._audit_matches())
    admin_page._export_audit()
    check("denetim CSV dosyasi yazildi", os.path.isfile(csv_path))
    import csv as _csv
    with open(csv_path, encoding="utf-8-sig", newline="") as fh:
        csv_rows = list(_csv.reader(fh, delimiter=";"))
    check("CSV basligi hash sutunlarini iceriyor",
          csv_rows[0][-2:] == ["onceki_ozet", "ozet"], str(csv_rows[0]))
    check("CSV tum satirlari yaziyor — ekrandaki 500 ile sinirli degil",
          len(csv_rows) - 1 == exported,
          "%d != %d" % (len(csv_rows) - 1, exported))
    check("disa aktarma denetim kaydina gecti",
          db.query_one("SELECT COUNT(*) AS n FROM audit_log"
                       " WHERE action = 'audit.export'")["n"] == 1)

    admin_page.audit_dates.combo.setCurrentIndex(
        admin_page.audit_dates.combo.findData("custom"))
    admin_page.audit_dates.start.setDate(old)
    admin_page.audit_dates.end.setDate(QtCore.QDate(2000, 12, 31))
    pump(app, 100)
    check("denetim kaydi tarihe gore suzuluyor",
          admin_page.audit_table.rowCount() == 0,
          "satir=%d" % admin_page.audit_table.rowCount())
    admin_page.audit_dates.combo.setCurrentIndex(0)
    pump(app, 100)

    made = backup.create(uid)
    check("arayuzden alinan yedek diskte", os.path.isfile(made))
    win._refresh_backup_label()
    check("durum cubugunda yedek yasi var",
          "Son yedek" in win.backup_label.text(), win.backup_label.text())

    # Notification center and jump-to-target
    win.home.refresh()
    items = notifications.collect(win.state.user)
    check("bildirim kutusu bildirim varken gorunur",
          win.home.notif_box.isHidden() != bool(items))
    check("bildirim listesi doldu",
          win.home.notif_list.count() == len(items),
          "liste=%d" % win.home.notif_list.count())
    if items:
        first = win.home.notif_list.item(0)
        check("bildirimde hedef sayfa saklaniyor",
              bool(first.data(Qt.UserRole)), str(first.data(Qt.UserRole)))
        win._go_target("approvals")
        check("bildirimden onay kuyruguna gidiliyor",
              win.tabs.currentWidget() is win.approvals)
        win._go_target("admin.perms")
        check("bildirimden yonetim bolumune gidiliyor",
              win.tabs.currentWidget() is win.admin
              and win.admin.tabs.currentIndex()
              == win.admin._sections["perms"])

    QtWidgets.QMessageBox.question = orig_question
    print("\n15. Olcum plani ve yeni ekranlar")
    from callog_common import i18n, points as points_svc, prefs, templates
    from callog_common import theme as theme_mod

    orig_question = QtWidgets.QMessageBox.question
    orig_information = QtWidgets.QMessageBox.information
    orig_warning = QtWidgets.QMessageBox.warning
    orig_get_text = QtWidgets.QInputDialog.getText
    orig_save = QtWidgets.QFileDialog.getSaveFileName
    orig_dir = QtWidgets.QFileDialog.getExistingDirectory
    QtWidgets.QMessageBox.question = staticmethod(
        lambda *a, **k: QtWidgets.QMessageBox.Yes)
    QtWidgets.QMessageBox.information = staticmethod(
        lambda *a, **k: QtWidgets.QMessageBox.Ok)
    QtWidgets.QMessageBox.warning = staticmethod(
        lambda *a, **k: QtWidgets.QMessageBox.Ok)

    plan_win = MainWindow(user)
    plan_win.show()
    pump(app, 200)
    psetup = plan_win.setup
    sim_id = db.query_one(
        "SELECT id FROM instruments WHERE driver = 'simulated'")["id"]
    psetup.instrument_combo.setCurrentIndex(
        psetup.instrument_combo.findData(sim_id))
    psetup.dut_company.setText("Plan Hastanesi")
    psetup.dut_manufacturer.setText("Fluke")
    psetup.dut_model.setText("175")
    psetup.dut_serial.setText("SN-PLAN-9")
    psetup.interval_spin.setValue(0.05)

    check("plan baslangicta bos", psetup.plan_table.rowCount() == 0)
    for fn, nominal, tol in (("VDC", 10.0, 0.05), ("VDC", 1.0, 0.02),
                             ("RES", 1000.0, 5.0)):
        psetup.function_combo.setCurrentIndex(
            psetup.function_combo.findData(fn))
        psetup.nominal_edit.setText(str(nominal))
        psetup.tolerance_edit.setText(str(tol))
        psetup._add_point()
    check("uc nokta plana eklendi", psetup.plan_table.rowCount() == 3)
    check("plan noktalari sozluk olarak okunuyor",
          [p["function"] for p in psetup.plan()] == ["VDC", "VDC", "RES"],
          str([p["function"] for p in psetup.plan()]))
    check("nokta birimi fonksiyondan geldi",
          psetup.plan()[2]["unit"] not in ("", None),
          str(psetup.plan()[2]))

    psetup.plan_table.selectRow(2)
    psetup._move_point(-1)
    check("nokta yukari tasindi",
          [p["nominal"] for p in psetup.plan()] == [10.0, 1000.0, 1.0],
          str([p["nominal"] for p in psetup.plan()]))
    psetup._move_point(1)
    psetup.plan_table.selectRow(2)
    psetup._remove_point()
    check("nokta kaldirildi", psetup.plan_table.rowCount() == 2)
    check("sira numaralari yeniden verildi",
          [psetup.plan_table.item(r, 0).text() for r in range(2)] == ["1", "2"])

    # Template: save, clear, reapply
    QtWidgets.QInputDialog.getText = staticmethod(
        lambda *a, **k: ("GUI plan sablonu", True))
    psetup._save_template()
    saved = templates.by_name("GUI plan sablonu")
    check("plan sablon olarak kaydedildi", saved is not None)
    check("sablonda iki nokta var", len(templates.points_of(saved)) == 2)
    check("sablon acilirda listelendi",
          psetup.template_combo.findData(saved["id"]) >= 0)

    psetup.plan_table.setRowCount(0)
    psetup._renumber_plan()
    psetup.template_combo.setCurrentIndex(
        psetup.template_combo.findData(saved["id"]))
    psetup._apply_template()
    check("sablon plana geri yuklendi", psetup.plan_table.rowCount() == 2)
    # 0.05 was requested but the box's lower limit is 0.1; the template
    # stores the value that ended up in the box, not the requested value.
    check("okuma periyodu sablondan geldi",
          abs(psetup.interval_spin.value() - 0.1) < 1e-6,
          str(psetup.interval_spin.value()))

    psetup.function_combo.setCurrentIndex(psetup.function_combo.findData("RES"))
    psetup.nominal_edit.setText("1000.0")
    psetup.tolerance_edit.setText("5.0")
    psetup._add_point()

    plan_started = {}
    psetup.session_started.connect(
        lambda sid, drv: plan_started.update({"sid": sid}))
    psetup._start()
    pump(app, 300)
    plan_sid = plan_started["sid"]
    plan_points = points_svc.list_for(plan_sid)
    check("plan veritabanina yazildi", len(plan_points) == 3)
    check("nokta sirasi korundu",
          [p["nominal"] for p in plan_points] == [10.0, 1.0, 1000.0],
          str([p["nominal"] for p in plan_points]))
    check("oturum sutunlari ilk noktayi tasiyor",
          db.query_one("SELECT nominal FROM sessions WHERE id = ?",
                       (plan_sid,))["nominal"] == 10.0)

    pacq = plan_win.acquire
    check("plan paneli coklu noktada gorunur", not pacq.plan_box.isHidden())
    check("sonraki nokta dugmesi acildi", not pacq.next_point_btn.isHidden())
    check("hedef satirinda nokta numarasi var",
          "1/3" in pacq.target_label.text(), pacq.target_label.text()[:60])

    for step in range(3):
        pacq._start_recording()
        pump(app, 500)
        check("%d. noktada okuma birikti" % (step + 1), pacq.stats.n > 2,
              "n=%d" % pacq.stats.n)
        if step < 2:
            before = pacq.point_index
            pacq._next_point()
            pump(app, 400)
            check("%d. noktaya gecildi" % (step + 2),
                  pacq.point_index == before + 1)
            check("yeni noktada istatistik sifirlandi", pacq.stats.n < 3,
                  "n=%d" % pacq.stats.n)
    check("simulasyon noktanin nominalini izliyor",
          abs(pacq.stats.mean - 1000.0) < 5.0, "mean=%.4f" % pacq.stats.mean)

    pacq.stop(status="completed")
    pump(app, 300)
    spread = db.query("SELECT point_id, COUNT(*) AS n FROM readings"
                      " WHERE session_id = ? GROUP BY point_id", (plan_sid,))
    check("okumalar uc noktaya dagildi", len(spread) == 3,
          str([(r["point_id"], r["n"]) for r in spread]))
    check("hicbir okuma sahipsiz kalmadi",
          all(r["point_id"] is not None for r in spread))
    check("butun noktalar tamamlandi",
          all(p["status"] == points_svc.DONE
              for p in points_svc.list_for(plan_sid)))

    plan_data = certificate.collect(plan_sid)
    check("sertifika uc nokta goruyor", len(plan_data["points"]) == 3)
    ppath, pcert_no, _pres = certificate.build_pdf(plan_sid, uid)
    check("coklu noktali sertifika yazildi",
          os.path.isfile(ppath) and os.path.getsize(ppath) > 20000,
          "%d bayt" % os.path.getsize(ppath))
    ptitles = [title for title, _rows
               in certificate.sections(plan_sid, pcert_no)[2]]
    check("sertifikada uc nokta bolumu var",
          sum(1 for x in ptitles if x.startswith("Ölçüm noktası")) == 3,
          str(ptitles))

    # Points in the history detail view
    phist = plan_win.history
    phist.reload()
    phist.focus_session(plan_sid)
    pump(app, 150)
    detail = phist.detail.toPlainText()
    check("gecmis detayinda nokta bloklari var",
          detail.count("nokta —") == 3, detail[:200])

    # --- General search ---------------------------------------------------
    from callog_common.ui.search_dialog import SearchDialog

    dlg = SearchDialog(plan_win)
    dlg.edit.setText("SN-PLAN-9")
    dlg._run()
    check("arama penceresi sonuc buluyor", dlg.list.count() > 1,
          "satir=%d" % dlg.list.count())
    picked = {}
    dlg.chosen.connect(lambda target, ident: picked.update(
        {"target": target, "id": ident}))
    for r in range(dlg.list.count()):
        item = dlg.list.item(r)
        if item.data(Qt.UserRole):
            dlg._activate(item)
            break
    check("secilen sonuc hedef bildiriyor", bool(picked.get("target")),
          str(picked))
    plan_win._open_search_hit(picked["target"], picked["id"])
    pump(app, 200)
    check("arama sonucundan sayfaya gidildi",
          plan_win.tabs.currentWidget() in (plan_win.history, plan_win.devices))
    check("kisa terimde sonuc yok", SearchDialog(plan_win).list.count() == 0)

    # --- Comparison --------------------------------------------------
    from callog_common.ui.compare_dialog import CompareDialog

    other_sid = db.query_one(
        "SELECT id FROM sessions WHERE id <> ? AND deleted_at IS NULL"
        "  AND status = 'completed' ORDER BY id LIMIT 1", (plan_sid,))["id"]
    cmp_dlg = CompareDialog([plan_sid, other_sid], plan_win)
    check("karsilastirma tablosu doldu", cmp_dlg.table.rowCount() >= 2,
          "satir=%d" % cmp_dlg.table.rowCount())
    check("karsilastirma aciklamasi seri sayisini yaziyor",
          "seri çizildi" in cmp_dlg.note.text(), cmp_dlg.note.text()[:80])
    check("birimi farkli nokta disarida birakiliyor",
          "Birimi farklı" in cmp_dlg.note.text()
          or cmp_dlg.table.rowCount() >= 2, cmp_dlg.note.text()[:120])

    # --- Batch operation ----------------------------------------------------
    batch_dir = os.path.join(_tmp, "toplu")
    os.makedirs(batch_dir, exist_ok=True)
    QtWidgets.QFileDialog.getExistingDirectory = staticmethod(
        lambda *a, **k: batch_dir)
    phist.reload()
    phist.table.clearSelection()
    for r in range(min(3, phist.table.rowCount())):
        phist.table.selectionModel().select(
            phist.table.model().index(r, 0),
            QtCore.QItemSelectionModel.Select | QtCore.QItemSelectionModel.Rows)
    pump(app, 100)
    chosen_ids = phist._selected_session_ids()
    check("coklu secim calisiyor", len(chosen_ids) >= 2, str(chosen_ids))
    check("karsilastir dugmesi coklu secimde acik",
          phist.compare_btn.isEnabled())
    phist._batch_excel()
    written = [f for f in os.listdir(batch_dir) if f.endswith(".xlsx")]
    check("toplu Excel dosyalari yazildi", len(written) == len(chosen_ids),
          "%d / %d" % (len(written), len(chosen_ids)))
    check("toplu aktarim denetim kaydina gecti",
          db.query_one("SELECT COUNT(*) AS n FROM audit_log"
                       " WHERE action = 'session.export_excel_batch'")["n"] == 1)

    before_certs = db.query_one(
        "SELECT COUNT(*) AS n FROM certificates")["n"]
    phist._batch_certificates()
    after_certs = db.query_one("SELECT COUNT(*) AS n FROM certificates")["n"]
    check("toplu sertifika ya uretti ya gerekcesini yazdi",
          after_certs >= before_certs, "%d -> %d" % (before_certs, after_certs))

    # --- Certificate preview ------------------------------------------
    from callog_common.ui import pdf_preview

    check("PDF onizleme modulu yuklenebiliyor", pdf_preview.is_available())
    preview = pdf_preview.PdfPreviewDialog(ppath, plan_win)
    check("onizleme belgeyi acti", preview.document.pageCount() >= 1,
          "sayfa=%d" % preview.document.pageCount())
    check("onizleme durum satiri dolu", "sayfa" in preview.status.text(),
          preview.status.text())
    preview.close()

    # --- Accessibility and language ----------------------------------------
    plan_win._set_scale(1.3)
    pump(app, 100)
    check("yazi olcegi uygulandi", abs(theme_mod.font_scale() - 1.3) < 0.01)
    check("olcek kullanici tercihine yazildi",
          abs(prefs.get_float(uid, prefs.FONT_SCALE, 1.0) - 1.3) < 0.01)
    # A "was it scaled" question rather than a concrete pixel value: keeps
    # the test from breaking if the base font size changes for design reasons.
    unscaled = theme_mod.stylesheet(theme_mod.current_mode())
    check("stil sayfasi buyutuldu",
          app.styleSheet() != unscaled
          and app.styleSheet() == theme_mod.scale_stylesheet(unscaled, 1.3),
          "olcek=%.2f" % theme_mod.font_scale())

    # --- Ctrl+/Ctrl- font-size shortcuts ---------------------------
    plan_win._set_scale(1.0)
    pump(app, 50)
    plan_win._zoom_in()
    pump(app, 50)
    check("ctrl+ yazi boyutunu buyuttu",
          abs(theme_mod.font_scale() - 1.1) < 0.01,
          "%.2f" % theme_mod.font_scale())
    plan_win._zoom_in()
    pump(app, 50)
    check("ikinci ctrl+ devam ediyor",
          abs(theme_mod.font_scale() - 1.2) < 0.01,
          "%.2f" % theme_mod.font_scale())
    plan_win._zoom_out()
    pump(app, 50)
    check("ctrl- yazi boyutunu kucultuyor",
          abs(theme_mod.font_scale() - 1.1) < 0.01,
          "%.2f" % theme_mod.font_scale())
    plan_win._zoom_reset()
    pump(app, 50)
    check("ctrl+0 yuzde 100e donuyor",
          abs(theme_mod.font_scale() - 1.0) < 0.01,
          "%.2f" % theme_mod.font_scale())

    # Must not overflow at the upper/lower limits
    plan_win._set_scale(theme_mod.MAX_SCALE)
    plan_win._zoom_in()
    pump(app, 50)
    check("ust sinirda tasmiyor",
          abs(theme_mod.font_scale() - theme_mod.MAX_SCALE) < 0.01,
          "%.2f" % theme_mod.font_scale())
    plan_win._set_scale(theme_mod.MIN_SCALE)
    plan_win._zoom_out()
    pump(app, 50)
    check("alt sinirda tasmiyor",
          abs(theme_mod.font_scale() - theme_mod.MIN_SCALE) < 0.01,
          "%.2f" % theme_mod.font_scale())
    plan_win._set_scale(1.0)
    pump(app, 50)

    # The actual menu shortcuts must also be wired up -- not just the
    # helper methods, but the key sequence the user will actually press
    zoom_actions = dict(
        (a.text(), a) for a in plan_win.menuBar().findChildren(QtGui.QAction)
        if a.text() in ("Yakınlaştır", "Uzaklaştır",
                        "Yazı boyutunu sıfırla (%100)"))
    check("yakinlastir Ctrl+= tusuna bagli",
          "Yakınlaştır" in zoom_actions
          and QtGui.QKeySequence("Ctrl+=") in zoom_actions["Yakınlaştır"].shortcuts(),
          str([s.toString() for s in
               zoom_actions.get("Yakınlaştır", QtGui.QAction()).shortcuts()]))
    check("sifirla Ctrl+0 tusuna bagli",
          "Yazı boyutunu sıfırla (%100)" in zoom_actions
          and zoom_actions["Yazı boyutunu sıfırla (%100)"].shortcut()
          == QtGui.QKeySequence("Ctrl+0"))

    plan_win._set_theme(theme_mod.CONTRAST)
    pump(app, 100)
    check("yuksek kontrast uygulandi",
          theme_mod.current_mode() == theme_mod.CONTRAST)
    check("kontrast menusu isaretlendi", plan_win.act_contrast.isChecked())
    plan_win._set_theme(theme_mod.LIGHT)
    plan_win._set_scale(1.0)

    plan_win._set_language(i18n.EN)
    check("dil ingilizceye alindi", i18n.language() == i18n.EN)
    check("dil tercihi kaydedildi", prefs.get(uid, prefs.LANGUAGE) == "en")
    en_win = MainWindow(user)
    labels = [en_win.tabs._pages[i][1].text()
              for i in range(en_win.tabs.count())]
    check("yeniden acilan pencerede serit ingilizce",
          any("History" in x for x in labels)
          and any("Devices" in x for x in labels), str(labels))
    en_win.close()
    plan_win._set_language(i18n.TR)
    check("turkceye donuldu", i18n.language() == i18n.TR)

    plan_win.close()
    QtWidgets.QMessageBox.question = orig_question
    QtWidgets.QMessageBox.information = orig_information
    QtWidgets.QMessageBox.warning = orig_warning
    QtWidgets.QInputDialog.getText = orig_get_text
    QtWidgets.QFileDialog.getSaveFileName = orig_save
    QtWidgets.QFileDialog.getExistingDirectory = orig_dir

    # --- 16. Layout: full screen with a 1920x1080 window ------------------
    print("\n16. Yerlesim")
    # Roughly what's left for the body on a 1920x1080 screen after the
    # taskbar, window frame, menu, and status bar are subtracted.
    MAX_W, MAX_H = 1900, 1000
    layout_win = MainWindow(user)
    layout_win.resize(1920, 1017)
    layout_win.show()
    pump(app, 300)

    wmin = layout_win.minimumSizeHint()
    check("pencere 1920x1080 ekrana sigiyor",
          wmin.width() <= MAX_W and wmin.height() <= MAX_H,
          "%d x %d" % (wmin.width(), wmin.height()))

    too_tall = []
    for i in range(layout_win.tabs.stack.count()):
        page = layout_win.tabs.stack.widget(i)
        if page.minimumSizeHint().height() > MAX_H:
            too_tall.append("%s=%d" % (type(page).__name__,
                                       page.minimumSizeHint().height()))
    check("hicbir sayfa ekrandan uzun degil", not too_tall, str(too_tall))

    check("uzun sayfa kaydirilabilir",
          layout_win.setup.minimumSizeHint().height() < 200,
          "min-h=%d" % layout_win.setup.minimumSizeHint().height())
    check("kaydirma alani kuruldu",
          bool(layout_win.setup.findChildren(QtWidgets.QScrollArea)))

    # Clipping watch: inputs and table rows must stay tall enough to fit the text.
    line_h = layout_win.setup.dut_company.sizeHint().height()
    text_h = layout_win.setup.dut_company.fontMetrics().height()
    check("giris kutusu metni kirpmiyor", line_h >= text_h + 6,
          "kutu=%d metin=%d" % (line_h, text_h))
    layout_win.devices.reload()
    pump(app, 150)
    row_h = layout_win.devices.table.verticalHeader().defaultSectionSize()
    check("tablo satiri metni kirpmiyor", row_h >= text_h + 4,
          "satir=%d metin=%d" % (row_h, text_h))
    check("tablo satiri gereksiz yuksek degil", row_h <= text_h + 12,
          "satir=%d metin=%d" % (row_h, text_h))

    # Accessibility: rows should also grow when the font is enlarged
    layout_win._set_scale(1.5)
    pump(app, 150)
    layout_win.devices.reload()
    pump(app, 150)
    big_text = layout_win.devices.table.fontMetrics().height()
    big_row = layout_win.devices.table.verticalHeader().defaultSectionSize()
    check("buyuk yazida satir da buyuyor", big_row >= big_text + 4,
          "satir=%d metin=%d" % (big_row, big_text))
    layout_win._set_scale(1.0)
    layout_win.close()

    print("\n" + "=" * 52)
    print("Gecen: %d    Kalan: %d" % (len(PASSED), len(FAILED)))
    if FAILED:
        print("\nBasarisiz:")
        for name in FAILED:
            print("  - %s" % name)
    print("=" * 52 + "\n")
    return 1 if FAILED else 0


if __name__ == "__main__":
    sys.exit(main())
