"""Certificate generation — PDF and DOCX.

Both formats are built from the same content (`sections()`). Once the
organization's official template arrives, only the layout will change; the
data collection, numbering, and approval logic stay the same.

The DOCX output is editable, so it's a practical stand-in until
organization-specific letterhead/signature fields are added.
"""

import hashlib
import os

from . import audit, branding, db, perms, points
from .i18n import t

#: The folder is computed **at call time**, not at import time: tests and
#: the screenshot script point `db.DATA_DIR` at a temporary folder. A fixed
#: module-level variable would miss that change and test output would land
#: in the real project folder.
def cert_dir():
    return os.path.join(db.DATA_DIR, "sertifikalar")

VERDICT_TR = {"pass": "UYGUN", "fail": "UYGUN DEĞİL", "info": "BİLGİLENDİRME AMAÇLI"}

CRITERION_TR = {
    "mean": "Ortalama ± U tolerans içinde",
    "minmax": "Tüm okumalar tolerans içinde",
}

SIM_WARNING = (
    "SİMÜLASYON ÇIKTISI — GEÇERLİ SERTİFİKA DEĞİLDİR\n"
    "Bu belge simüle edilmiş bir cihazdan üretilmiştir. İçindeki ölçüm "
    "değerleri gerçek bir kalibrasyona ait değildir; yalnızca deneme ve eğitim "
    "amacıyla kullanılabilir."
)

UNCERTAINTY_NOTE = (
    "Beyan edilen genişletilmiş belirsizlik, standart belirsizliğin k=2 kapsam "
    "çarpanı ile çarpılmasıyla elde edilmiştir ve yaklaşık %95 kapsam olasılığına "
    "karşılık gelir. Bu sürümde yalnızca A tipi belirsizlik bileşeni "
    "hesaplanmaktadır; B tipi bileşenler eklendiğinde değer büyüyecektir."
)


def next_cert_no(simulated=False):
    """Generates a sequential number in CAL-MED-YYYY-NNNN form.

    Simulation output gets a number from a separate series (SIM-CAL-MED-...),
    so documents produced for testing don't consume official certificate
    numbers and can be told apart just by looking at the number.
    """
    from datetime import datetime

    year = datetime.now().year
    prefix = "%sCAL-MED-%d-" % ("SIM-" if simulated else "", year)
    rows = db.query(
        "SELECT cert_no FROM certificates WHERE cert_no LIKE ?", (prefix + "%",))
    numbers = [int(r["cert_no"].rsplit("-", 1)[1]) for r in rows
               if r["cert_no"].startswith(prefix)]
    return "%s%04d" % (prefix, max(numbers) + 1 if numbers else 1)


def for_session(session_id):
    """The session's certificate; None if there isn't one."""
    return db.query_one(
        "SELECT * FROM certificates WHERE session_id = ?", (session_id,))


def for_series(series_id):
    """The waveform series measurement's certificate; None if there isn't one."""
    if not series_id:
        return None
    return db.query_one(
        "SELECT * FROM certificates WHERE series_id = ?", (series_id,))


#: Shared mapping that links a certificate to its device.
#:
#: A certificate can come from two sources (measurement session / waveform
#: series measurement), and the device link is built differently for each:
#: `sessions.dut_id` versus `waveform_captures.dut_id`. A single plain JOIN
#: would drop series certificates from listings entirely. This is kept as
#: one shared piece of SQL so every listing screen uses the same mapping —
#: otherwise two screens could disagree on "which certificates exist".
SOURCE_JOIN = (
    " JOIN ("
    "   SELECT c2.id AS cid, s.dut_id AS dut_id, 'session' AS kind"
    "     FROM certificates c2"
    "     JOIN sessions s ON s.id = c2.session_id"
    "                    AND s.deleted_at IS NULL"
    "   UNION ALL"
    "   SELECT c2.id, (SELECT w.dut_id FROM waveform_captures w"
    "                   WHERE w.series_id = c2.series_id LIMIT 1),"
    "          'series'"
    "     FROM certificates c2 WHERE c2.series_id IS NOT NULL"
    " ) k ON k.cid = c.id"
)

#: Turkish label for the certificate type — the "Tür" column in listings.
KIND_TR = {"session": "ölçüm oturumu", "series": "dalga serisi"}


def pending(limit=200):
    """Certificates awaiting approval — **oldest first**.

    The order is deliberately reversed: the queue starts with the
    longest-waiting document. Sorted newest-first, the oldest document
    would sink to the bottom of the list during a busy week and get
    forgotten there.
    """
    return db.query(
        "SELECT c.*, k.dut_id, k.kind, d.manufacturer, d.model, d.serial_no,"
        " d.company, u.full_name AS issued_by_name"
        " FROM certificates c" + SOURCE_JOIN +
        # LEFT JOIN is required: a waveform series can be captured without
        # being linked to any device (when the dut_combo has "(baglama)"
        # selected), in which case k.dut_id is NULL. With an INNER JOIN,
        # `d.id = NULL` would never match and the certificate would
        # silently drop out of the queue even while awaiting approval.
        " LEFT JOIN duts d ON d.id = k.dut_id"
        " JOIN users u ON u.id = c.issued_by"
        " WHERE c.approved_at IS NULL AND c.deleted_at IS NULL"
        " ORDER BY c.issued_at, c.id LIMIT ?", (int(limit),))


def device_label(row):
    """Device label for the certificate.

    If a waveform series was captured without being linked to a DUT,
    `manufacturer` comes back NULL (see the note on `pending()`) — in that
    case, instead of showing "None None — None", it explicitly shows
    "cihaza bağlanmadı".
    """
    if row["manufacturer"] is None:
        return "cihaza bağlanmadı"
    return "%s %s — %s" % (row["manufacturer"], row["model"], row["serial_no"])


def register_series(series_id, cert_no, issued_by, result, pdf_path,
                    pdf_sha256):
    """Records the series shock report in the certificate register.

    The report number already lives in `waveform_captures.report_no`, but
    that field only answers "was it produced?" — no approval, soft
    deletion, or certificate listing. If the report is regenerated for the
    same series, the **existing row is updated** — inserting a new row
    would violate the `series_id` uniqueness constraint and make the series
    look certified twice. Approval info is deliberately reset: since the
    content changed, the old approval no longer applies to this document.

    Returns: the certificate record's id.
    """
    existing = for_series(series_id)
    if existing is not None:
        db.execute(
            "UPDATE certificates SET cert_no = ?, issued_at = ?, issued_by = ?,"
            " result = ?, pdf_path = ?, pdf_sha256 = ?, approved_by = NULL,"
            " approved_at = NULL WHERE id = ?",
            (cert_no, db.utc_now(), issued_by, result, pdf_path, pdf_sha256,
             existing["id"]))
        cid = existing["id"]
        action = "certificate.reissue"
    else:
        cid = db.execute(
            "INSERT INTO certificates (series_id, cert_no, issued_at, issued_by,"
            " result, pdf_path, pdf_sha256) VALUES (?,?,?,?,?,?,?)",
            (series_id, cert_no, db.utc_now(), issued_by, result, pdf_path,
             pdf_sha256))
        action = "certificate.issue"
    audit.log(action, user_id=issued_by, entity="certificate", entity_id=cid,
              detail={"cert_no": cert_no, "series_id": series_id,
                      "result": result, "sha256": pdf_sha256,
                      "kind": "waveform_series"})
    return cid


def collect(session_id):
    """Gathers all data needed for the certificate into a single dict.

    Measurements are now computed point by point (`points.collect`). The
    top-level `n`, `mean`, `U`, … fields in the dict are the **first
    point's** numbers: for a single-point session that matches existing
    behavior, and for a multi-point session it's a summary that fits into
    the listing's narrow columns. All points are in ``d["points"]``.

    `result`, however, is the decision for the **whole session**: if even
    one point isn't within tolerance, the document isn't either.
    """
    s = db.query_one("SELECT * FROM sessions WHERE id = ?", (session_id,))
    if s is None:
        raise ValueError("Oturum bulunamadı: %s" % session_id)

    dut = db.query_one("SELECT * FROM duts WHERE id = ?", (s["dut_id"],))
    inst = db.query_one("SELECT * FROM instruments WHERE id = ?", (s["instrument_id"],))
    op = db.query_one("SELECT * FROM users WHERE id = ?", (s["operator_id"],))

    summaries = points.collect(session_id)
    first = summaries[0]

    data = {
        "session": s, "dut": dut, "instrument": inst, "operator": op,
        "points": summaries, "multi": len(summaries) > 1,
        "total_n": sum(p["n"] for p in summaries),
        "result": points.overall_result(summaries),
        "point_result": first["result"],
    }
    # Backward-compatible fields — the first point's numbers
    for key in ("n", "excluded", "mean", "std", "u_a", "U", "nominal",
                "tolerance", "mode", "deviation", "min", "max", "unit",
                "function"):
        data[key] = first[key]
    return data


def sections(session_id, cert_no=None):
    """Returns the document's content in a format-independent structure.

    Returns: (cert_no, is_simulated, [(section_name, [(label, value), ...]), ...],
            result_text, footnote, signature_rows)
    """
    d = collect(session_id)
    s, dut, inst, op = d["session"], d["dut"], d["instrument"], d["operator"]
    unit = s["unit"]

    def fmt(v):
        return "—" if v is None else "%.6g %s" % (v, unit)

    simulated = bool(s["is_simulated"])
    cert_no = cert_no or next_cert_no(simulated=simulated)

    tol_text = "—"
    if d["tolerance"]:
        tol_text = "± %g %s   (%g … %g %s)" % (
            d["tolerance"], unit,
            d["nominal"] - d["tolerance"], d["nominal"] + d["tolerance"], unit)

    body = [
        ("", [
            ("Sertifika no", cert_no),
            ("Veriliş tarihi", db.utc_now()[:10]),
            ("Ölçüm tarihi", (s["started_at"] or "")[:10]),
            ("Ölçüm oturumu", (s["name"] or "").strip() or ("#%d" % s["id"])),
        ]),
        ("Kalibre edilen cihaz", [
            ("Şirket / müşteri", dut["company"]),
            ("Üretici firma", dut["manufacturer"]),
            ("Model", dut["model"]),
            ("Seri no", dut["serial_no"]),
            ("Cihaz tipi", dut["device_type"] or "—"),
        ]),
        ("Kullanılan referans standart", [
            ("Cihaz", "%s %s" % (inst["brand"], inst["model"])),
            ("Seri no", inst["serial_no"]),
            ("Kalibrasyon sertifika no", inst["cal_cert_no"] or "—"),
            ("Geçerlilik", inst["cal_due"] or "—"),
        ]),
        ("Ortam şartları", [
            ("Sıcaklık / nem / basınç", "%s °C    %s %%RH    %s kPa    (%s)" % (
                s["env_temp"] if s["env_temp"] is not None else "—",
                s["env_rh"] if s["env_rh"] is not None else "—",
                s["env_pressure"] if s["env_pressure"] is not None else "—",
                t("elle girildi") if s["env_source"] == "manual"
                else (s["env_source"] or "—"))),
        ]),
    ]

    # A single-point session stays in its existing single-section form. In
    # a multi-point session, each point gets **its own section**: each
    # point has its own tolerance, criterion, and uncertainty, and cramming
    # them all into one table would make the certificate unreadable and
    # hard to trace. The plan summary above shows at a glance where each
    # point stands.
    if d["multi"]:
        body.append(("Ölçüm planı", [
            ("%d. %s" % (p["seq"], points.label(p)),
             t("n = %d   ·   %s") % (p["n"], t(VERDICT_TR[p["result"]])))
            for p in d["points"]]))
        for p in d["points"]:
            body.append((t("Ölçüm noktası %d — %s")
                         % (p["seq"], points.label(p)),
                         _point_rows(p, with_result=True)))
    else:
        body.append(("Ölçüm sonuçları", _measurement_rows(s, d, fmt, tol_text)))

    signatures = [
        ("Ölçümü yapan", op["full_name"]),
        ("Onaylayan", "................................................"),
    ]
    return (cert_no, simulated, body, t(VERDICT_TR[d["result"]]),
            t(UNCERTAINTY_NOTE), signatures, d)


def _measurement_rows(s, d, fmt, tol_text):
    rows = [("Ölçüm fonksiyonu", s["function"])]
    # Which channel was measured on the oscilloscope is part of the result:
    # a "Vpp = 1.984 V" row without the channel isn't traceable on its own.
    channel = _channel_of(s)
    if channel:
        rows.append(("Ölçülen kanal", channel))
    rows += [
            ("Okuma sayısı (n)",
         t("%d  (dışlanan: %d)") % (d["n"], d["excluded"])),
            ("Nominal değer", fmt(d["nominal"])),
            ("Tolerans", tol_text),
            ("Uygunluk kriteri", t(CRITERION_TR[d["mode"]])),
            ("Ölçülen ortalama", fmt(d["mean"])),
            ("Standart sapma (s)", fmt(d["std"])),
            ("En küçük / en büyük", "%s / %s" % (fmt(d["min"]), fmt(d["max"]))),
            ("Sapma", fmt(d["deviation"])),
            ("Genişletilmiş belirsizlik U (k=2)", fmt(d["U"])),
    ]
    return rows


def _point_rows(p, with_result=False):
    """Rows for a single measurement point.

    The unit comes from the point itself: the same session could measure
    10 V and 1 kΩ, and using the session's unit would print the latter as
    "1000 V".
    """
    unit = p["unit"]

    def fmt(v):
        return "—" if v is None else "%.6g %s" % (v, unit)

    tol_text = "—"
    if p["tolerance"]:
        tol_text = "± %g %s   (%g … %g %s)" % (
            p["tolerance"], unit,
            p["nominal"] - p["tolerance"], p["nominal"] + p["tolerance"], unit)

    rows = [("Ölçüm fonksiyonu", p["function"])]
    channel = (p["channel"] or "").strip()
    if channel:
        rows.append(("Ölçülen kanal", channel.replace("CHANnel", "Kanal ")))
    rows += [
        ("Okuma sayısı (n)",
         t("%d  (dışlanan: %d)") % (p["n"], p["excluded"])),
        ("Nominal değer", fmt(p["nominal"])),
        ("Tolerans", tol_text),
        ("Uygunluk kriteri", t(CRITERION_TR[p["mode"]])),
        ("Ölçülen ortalama", fmt(p["mean"])),
        ("Standart sapma (s)", fmt(p["std"])),
        ("En küçük / en büyük", "%s / %s" % (fmt(p["min"]), fmt(p["max"]))),
        ("Sapma", fmt(p["deviation"])),
        ("Genişletilmiş belirsizlik U (k=2)", fmt(p["U"])),
    ]
    if with_result:
        rows.append(("Sonuç", t(VERDICT_TR[p["result"]])))
    return rows


def _channel_of(session_row):
    """Whether the session has a channel recorded — older databases lack the column."""
    try:
        keys = session_row.keys()
    except AttributeError:
        return None
    if "channel" not in keys:
        return None
    value = (session_row["channel"] or "").strip()
    if not value:
        return None
    return value.replace("CHANnel", "Kanal ")


# --- PDF ------------------------------------------------------------------
def build_pdf(session_id, issued_by, cert_no=None):
    """Generates the certificate as a PDF, saves it to the database, returns its path."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                    TableStyle)

    from . import pdffont

    font, font_bold, ascii_only = pdffont.register()
    cert_no, simulated, body, verdict, note, signatures, data = sections(
        session_id, cert_no)

    directory = cert_dir()
    if not os.path.isdir(directory):
        os.makedirs(directory)
    path = os.path.join(directory, "%s.pdf" % cert_no)

    base = getSampleStyleSheet()
    p_body = ParagraphStyle("govde", parent=base["BodyText"], fontName=font,
                            fontSize=9, leading=13)
    p_h1 = ParagraphStyle("baslik", parent=base["Title"], fontName=font_bold,
                          fontSize=15, leading=19, spaceAfter=2)
    p_h2 = ParagraphStyle("altbaslik", parent=base["Heading2"], fontName=font_bold,
                          fontSize=11.5, leading=15, spaceBefore=2)
    p_h4 = ParagraphStyle("bolum", parent=base["Heading4"], fontName=font_bold,
                          fontSize=9.5, leading=13, spaceBefore=4, spaceAfter=2)

    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=cert_no, author=branding.org_name(),
    )

    def kv_table(rows):
        t = Table([[k, v] for k, v in rows], colWidths=(58 * mm, 102 * mm))
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), font),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#555555")),
            ("LINEBELOW", (0, 0), (-1, -2), 0.25, colors.HexColor("#DDDDDD")),
        ]))
        return t

    story = [
        Paragraph(branding.header_line(), p_h1),
        Paragraph(t("KALİBRASYON SERTİFİKASI"), p_h2),
        Spacer(1, 5 * mm),
    ]

    if simulated:
        head, rest = t(SIM_WARNING).split("\n", 1)
        story += [
            Table([[Paragraph("<b>%s</b><br/>%s" % (head, rest), p_body)]],
                  colWidths=[160 * mm],
                  style=TableStyle([
                      ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FCEBEB")),
                      ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#A32D2D")),
                      ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#A32D2D")),
                      ("LEFTPADDING", (0, 0), (-1, -1), 8),
                      ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                      ("TOPPADDING", (0, 0), (-1, -1), 6),
                      ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                  ])),
            Spacer(1, 5 * mm),
        ]

    # Section titles and row labels are translated from one single spot:
    # filtering them through here instead of wrapping every string inside
    # `sections()` individually removes the risk of forgetting the
    # translation when a new row is added.
    for title, rows in body:
        if title:
            story.append(Paragraph(t(title), p_h4))
        story.append(kv_table([(t(k), v) for k, v in rows]))
        story.append(Spacer(1, 4 * mm))

    # --- measurement charts (one per point) --------------------------------
    from . import chart

    for i, p in enumerate(data["points"]):
        drawing = chart.session_drawing(session_id, font=font,
                                        font_bold=font_bold, summary=p,
                                        is_first=(i == 0))
        if drawing is None:
            continue
        title = (t("Ölçüm grafiği") if not data["multi"]
                 else t("Ölçüm grafiği — %d. nokta, %s")
                 % (p["seq"], points.label(p)))
        story += [Paragraph(title, p_h4), drawing, Spacer(1, 4 * mm)]

    story += [
        Paragraph(t("Sonuç: %s") % verdict, p_h2),
        Spacer(1, 3 * mm),
        Paragraph(note, p_body),
        Spacer(1, 9 * mm),
        kv_table([(t(k), v) for k, v in signatures]),
    ]

    def stamp(canvas, _doc):
        """Stamps a diagonal watermark onto simulation output.

        There's also a warning banner above, but the watermark still shows
        up on a photocopy or a screenshot — that's what actually keeps the
        document from being mistaken for a real certificate.
        """
        if not simulated:
            return
        canvas.saveState()
        canvas.translate(A4[0] / 2.0, A4[1] / 2.0)
        canvas.rotate(45)
        canvas.setFillColor(colors.Color(0.64, 0.18, 0.18, alpha=0.16))
        canvas.setFont(font_bold, 62)
        canvas.drawCentredString(0, 0, t("SİMÜLASYON"))
        canvas.setFont(font_bold, 20)
        canvas.drawCentredString(0, -42, t("GEÇERLİ SERTİFİKA DEĞİLDİR"))
        canvas.restoreState()

    doc.build(story, onFirstPage=stamp, onLaterPages=stamp)

    with open(path, "rb") as fh:
        sha = hashlib.sha256(fh.read()).hexdigest()

    cid = db.execute(
        "INSERT INTO certificates (session_id, cert_no, issued_at, issued_by,"
        " result, pdf_path, pdf_sha256) VALUES (?,?,?,?,?,?,?)",
        (session_id, cert_no, db.utc_now(), issued_by, data["result"], path, sha),
    )
    audit.log("certificate.issue", user_id=issued_by, entity="certificate",
              entity_id=cid, detail={"cert_no": cert_no, "sha256": sha,
                                     "result": data["result"],
                                     "simulated": simulated,
                                     "font_ascii_fallback": ascii_only})
    return path, cert_no, data["result"]


# --- DOCX -----------------------------------------------------------------
def write_docx(session_id, path, cert_no=None):
    """Writes the same content as an editable Word document.

    Doesn't create a certificate record — the PDF is the official output;
    the DOCX is provided for working on it further (organization
    letterhead, an extra note, a signature). If the certificate hasn't been
    issued yet, the number field shows the next number.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    existing = db.query_one(
        "SELECT cert_no FROM certificates WHERE session_id = ?", (session_id,))
    cert_no, simulated, body, verdict, note, signatures, _data = sections(
        session_id, cert_no or (existing["cert_no"] if existing else None))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(branding.header_line())
    run.bold = True
    run.font.size = Pt(15)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(t("KALİBRASYON SERTİFİKASI"))
    run.bold = True
    run.font.size = Pt(12)

    if simulated:
        warn = doc.add_paragraph()
        head, rest = t(SIM_WARNING).split("\n", 1)
        run = warn.add_run(head)
        run.bold = True
        run.font.color.rgb = RGBColor(0xA3, 0x2D, 0x2D)
        run = warn.add_run("\n" + rest)
        run.font.color.rgb = RGBColor(0xA3, 0x2D, 0x2D)

    for section_title, rows in body:
        if section_title:
            heading = doc.add_paragraph()
            run = heading.add_run(t(section_title))
            run.bold = True
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = str(t(label))
            cells[1].text = str(value)
            cells[0].paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

    # Measurement chart — PNG counterpart of the vector drawing in the PDF
    from . import chart, pdffont

    font_name, font_bold, _ = pdffont.register()
    chart_png = os.path.join(
        os.path.dirname(os.path.abspath(path)) or ".",
        ".%s-grafik.png" % os.path.splitext(os.path.basename(path))[0])
    try:
        if chart.png(session_id, chart_png, font=font_name,
                     font_bold=font_bold):
            heading = doc.add_paragraph()
            heading.add_run(t("Ölçüm grafiği")).bold = True
            doc.add_picture(chart_png, width=Inches(6.3))
            os.remove(chart_png)
    except Exception:
        # The document should still come out complete if the chart fails
        pass

    result = doc.add_paragraph()
    run = result.add_run(t("Sonuç: %s") % verdict)
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(note)
    doc.add_paragraph()

    sig = doc.add_table(rows=0, cols=2)
    sig.style = "Table Grid"
    for label, value in signatures:
        cells = sig.add_row().cells
        cells[0].text = str(t(label))
        cells[1].text = str(value)
        cells[0].paragraphs[0].runs[0].bold = True

    doc.save(path)
    audit.log("certificate.export_docx", entity="session", entity_id=session_id,
              detail={"cert_no": cert_no, "path": path})
    return path, cert_no


# --- approval and deletion -------------------------------------------------
def approve(cert_id, approver_id):
    """Approval by the lab manager. The certificate is locked after approval."""
    perms.require_actor(approver_id, perms.CERT_APPROVE)
    row = db.query_one("SELECT * FROM certificates WHERE id = ?", (cert_id,))
    if row is None:
        raise ValueError("Sertifika bulunamadı")
    if row["deleted_at"]:
        raise ValueError("Silinmiş bir sertifika onaylanamaz")
    if row["approved_at"]:
        raise ValueError("Bu sertifika zaten onaylanmış")

    db.execute(
        "UPDATE certificates SET approved_by = ?, approved_at = ? WHERE id = ?",
        (approver_id, db.utc_now(), cert_id),
    )
    audit.log("certificate.approve", user_id=approver_id, entity="certificate",
              entity_id=cert_id, detail={"cert_no": row["cert_no"]})


def soft_delete(cert_id, user_id, reason):
    """Marks the certificate as deleted.

    The record doesn't leave the database: no gap forms in the number
    series, and the measurement data and audit trail are preserved. Only
    admins see deleted records.
    """
    perms.require_actor(user_id, perms.CERT_DELETE)
    row = db.query_one("SELECT * FROM certificates WHERE id = ?", (cert_id,))
    if row is None:
        raise ValueError("Sertifika bulunamadı")
    if row["deleted_at"]:
        raise ValueError("Bu sertifika zaten silinmiş")
    if not (reason or "").strip():
        raise ValueError("Silme gerekçesi zorunludur")

    db.execute(
        "UPDATE certificates SET deleted_at = ?, deleted_by = ?, delete_reason = ?"
        " WHERE id = ?", (db.utc_now(), user_id, reason.strip(), cert_id))
    audit.log("certificate.delete", user_id=user_id, entity="certificate",
              entity_id=cert_id,
              detail={"cert_no": row["cert_no"], "reason": reason.strip()})


def restore(cert_id, user_id):
    """Restores a deleted certificate (admin only)."""
    perms.require_actor(user_id, perms.CERT_RESTORE)
    row = db.query_one("SELECT * FROM certificates WHERE id = ?", (cert_id,))
    if row is None or not row["deleted_at"]:
        raise ValueError("Silinmiş bir sertifika değil")
    db.execute(
        "UPDATE certificates SET deleted_at = NULL, deleted_by = NULL,"
        " delete_reason = NULL WHERE id = ?", (cert_id,))
    audit.log("certificate.restore", user_id=user_id, entity="certificate",
              entity_id=cert_id, detail={"cert_no": row["cert_no"]})
